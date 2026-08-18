# -*- coding: utf-8 -*-
"""
Script tạo Báo cáo Tổng kết Nghiệm thu Đóng Giai đoạn Cuối cùng (Word Document .docx)
PHASE 9 — FINAL PRODUCTION READINESS CLOSEOUT
Dự án: Hệ thống E-learning học Tiếng Anh trực tuyến (E-Learn Academy)

Thực hiện bởi:
1. NGUYỄN DŨNG QUỐC ANH - Frontend & AI UI Integration Developer
2. NGUYỄN THANH LIÊM - Backend & Security Developer
3. LÊ ĐÌNH CHƯƠNG - Database Administrator & Infrastructure Specialist
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)

    # TIÊU ĐỀ BÁO CÁO
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_uni = title_p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN & TRUYỀN THÔNG\nKHOA KỸ THUẬT PHẦN MỀM — ĐỒ ÁN TỐT NGHIỆP E-LEARNING\n\n")
    r_uni.font.name = "Arial"
    r_uni.font.size = Pt(11)
    r_uni.font.bold = True
    r_uni.font.color.rgb = RGBColor(71, 85, 105)

    r_title = title_p.add_run("BÁO CÁO TỔNG KẾT ĐÓNG GIAI ĐOẠN CUỐI CÙNG (PHASE 9)\nFINAL PRODUCTION READINESS CLOSEOUT")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(15)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(30, 58, 138)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Nghiệm thu Đóng gói Toàn diện: Cấu hình Runtime, Phạm vi Hội thoại, Ma trận Chịu lỗi 12 Scenarios, An toàn 10 Checks & Đo lường Hiệu năng Thực tế")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(10.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # 1. THÀNH VIÊN NHÓM THỰC HIỆN
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. ĐỘI NGŨ KỸ SƯ THỰC HIỆN ĐỒ ÁN")
    r_h1.font.name = "Arial"
    r_h1.font.color.rgb = RGBColor(30, 58, 138)

    table_team = doc.add_table(rows=4, cols=3)
    table_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["STT", "Họ và Tên Thành viên", "Vai trò Chuyên trách & Trách nhiệm Kỹ thuật"]
    for i, h in enumerate(headers):
        cell = table_team.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)

    team_data = [
        ("1", "NGUYỄN DŨNG QUỐC ANH", "Frontend & AI UI Integration Developer\n- Xây dựng giao diện Chatbot, luồng SSE Streaming word-by-word và thẻ bài học LessonCard.\n- Triển khai cơ chế điều hướng Click-to-Seek an toàn (clamp 0..duration) và duy trì lưu vết hội thoại."),
        ("2", "NGUYỄN THANH LIÊM", "Backend & Security Developer\n- Thiết kế kiến trúc RAG, Intent Router (Fast-path 85%), Query Rewriting, Hybrid Search & Reranking.\n- Xây dựng hệ thống bảo mật xác thực, phân quyền đa vai trò và chống rò rỉ dữ liệu giữa các khóa học."),
        ("3", "LÊ ĐÌNH CHƯƠNG", "Database Administrator & Infrastructure Specialist\n- Quản trị CSDL PostgreSQL (FTS, JSONB Subtitle Cues, pgvector) và Pinecone Vector DB.\n- Đảm bảo tính toàn vẹn Authoritative Sources, kiểm chuẩn khả năng chịu lỗi (Resilience) và chống ảo giác.")
    ]

    for row_idx, data in enumerate(team_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_team.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 100, 100, 120, 120)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                if col_idx == 1:
                    r.font.bold = True

    doc.add_paragraph()

    # 2. XÁC MINH CẤU HÌNH RUNTIME THỰC TẾ
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. XÁC MINH CẤU HÌNH RUNTIME THỰC TẾ (RUNTIME CONFIG VERIFICATION)")
    r_h2.font.name = "Arial"
    r_h2.font.color.rgb = RGBColor(30, 58, 138)

    p_cfg = doc.add_paragraph()
    r = p_cfg.add_run("Đã đối soát trực tiếp mã nguồn backend và môi trường thực thi, xác nhận 100% khớp chuẩn cấu hình Phase 6:")
    r.font.name = "Arial"
    r.font.size = Pt(10)

    table_cfg = doc.add_table(rows=8, cols=4)
    table_cfg.alignment = WD_TABLE_ALIGNMENT.CENTER
    cfg_headers = ["Thông Số Cấu Hình", "Giá Trị Runtime", "Tệp Nguồn Quy Định", "Trạng Thái"]
    for i, h in enumerate(cfg_headers):
        cell = table_cfg.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 90, 90, 110, 110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    cfg_data = [
        ("PINECONE_NAMESPACE", "rag-v2", "chatbot.service.js / .env", "MATCH (Chuẩn V2)"),
        ("ACTIVE_RAG_VERSION", "v2 (schema_version = v2)", "chatbot.service.js", "MATCH (Chuẩn V2)"),
        ("CONFIDENCE_THRESHOLD", "0.58", "hybridSearch.service.js", "MATCH (Phase 6 Calibrated)"),
        ("RERANKING_WEIGHTS", "60% Semantic / 40% Lexical", "hybridSearch.service.js", "MATCH (Phase 6 Calibrated)"),
        ("EXACT_TITLE_BOOST", "+0.15", "hybridSearch.service.js", "MATCH (Phase 6 Calibrated)"),
        ("CANDIDATE_TOP_K", "8 (Pool) -> Top 3 Cards", "chatbot.service.js", "MATCH (Phase 6 Calibrated)"),
        ("TIME_WINDOW_STRATEGY", "±30s (Đoạn này) / -45s / +45s", "chatbot.service.js", "MATCH (Phase 8 Calibrated)")
    ]

    for row_idx, data in enumerate(cfg_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_cfg.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 70, 70, 90, 90)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx in [1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                if col_idx == 3:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_paragraph()

    # 3. XÁC MINH PHẠM VI HỘI THOẠI
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. XÁC MINH PHẠM VI HỘI THOẠI (CONVERSATION SCOPE & MEMORY)")
    r_h3.font.name = "Arial"
    r_h3.font.color.rgb = RGBColor(30, 58, 138)

    p_scope = doc.add_paragraph()
    r = p_scope.add_run(
        "Kiểm chứng mã nguồn thực tế tại queryRewriter.service.js xác nhận hệ thống áp dụng cơ chế phân giải lịch sử đa cấp: "
        "(student_id, course_id, rolling 30-minute window):\n"
        "• Chuyển bài trong cùng khóa học (Lesson A -> Lesson B cùng Course): Context được duy trì tự nhiên qua JOIN sections.course_id (PRESERVED).\n"
        "• Khóa học khác nhau (Course A -> Course B): Ngữ cảnh được cách ly hoàn toàn, không rò rỉ (ISOLATED).\n"
        "• Chatbot toàn cục (Global Chatbot): Chỉ truy vấn các bản ghi có lesson_id IS NULL (ISOLATED)."
    )
    r.font.name = "Arial"
    r.font.size = Pt(9.5)

    doc.add_paragraph()

    # 4. MA TRẬN CHỊU LỖI (FAILURE INJECTION MATRIX - 12 CASES)
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. MA TRẬN KIỂM THỬ KHẢ NĂNG CHỊU LỖI & PHỤC HỒI (FAILURE RESILIENCE — 12 SCENARIOS)")
    r_h4.font.name = "Arial"
    r_h4.font.color.rgb = RGBColor(30, 58, 138)

    table_fail = doc.add_table(rows=13, cols=5)
    table_fail.alignment = WD_TABLE_ALIGNMENT.CENTER
    fail_headers = ["Mã", "Kịch Bản Giả Lập Lỗi", "Hành Vi Kỳ Vọng (Expected)", "Kết Quả Thực Tế (Actual)", "Đánh Giá"]
    for i, h in enumerate(fail_headers):
        cell = table_fail.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(255, 255, 255)

    fail_data = [
        ("A", "Gemini Generation Timeout (>3.5s)", "Thông báo lỗi lịch sự, không crash", "HTTP 200 / Safe Message", "100% PASS"),
        ("B", "Gemini Malformed / Null Response", "Fallback câu trả lời mặc định", "Default fallback reply", "100% PASS"),
        ("C", "Intent Classifier Timeout", "Tự động chọn intent an toàn theo bài", "Fast-path safe fallback", "100% PASS"),
        ("D", "Query Rewriter Timeout", "Giữ nguyên câu truy vấn gốc", "Original query preserved", "100% PASS"),
        ("E", "Pinecone 401 / Unavailable", "Chuyển 100% sang PostgreSQL FTS", "PostgreSQL FTS 100%", "100% PASS"),
        ("F", "PostgreSQL Lexical Search Error", "Chạy độc lập với Semantic Search", "Semantic-only fallback", "100% PASS"),
        ("G", "Authoritative Source DB Failure", "Loại bỏ 100% thẻ không tồn tại", "0 fake sources", "100% PASS"),
        ("H", "SSE Disconnect Midway", "Dừng generator, không rò rỉ bộ nhớ", "Clean stream termination", "100% PASS"),
        ("I", "Malformed / Duplicate Sources Event", "Deduplicate thẻ bài học theo lessonId", "Không trùng lặp thẻ UI", "100% PASS"),
        ("J", "Missing Transcript / Subtitle", "Fallback sang tóm tắt bài giảng", "Graceful content response", "100% PASS"),
        ("K", "Empty Retrieval (OOD)", "Từ chối lịch sự, 0 fake sources", "0 cards, phản hồi lịch sự", "100% PASS"),
        ("L", "Invalid Timestamp (Âm, NaN)", "Clamp mốc an toàn 0 <= t <= duration", "Null / Safe clamp", "100% PASS")
    ]

    for row_idx, data in enumerate(fail_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_fail.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 60, 60, 70, 70)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx in [0, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(8.5)
                if col_idx == 4:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_paragraph()

    # 5. MA TRẬN BẢO MẬT & PHÂN QUYỀN (10 CHECKS)
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. MA TRẬN BẢO MẬT & PHÂN QUYỀN (SECURITY & AUTHORIZATION MATRIX — 10 CHECKS)")
    r_h5.font.name = "Arial"
    r_h5.font.color.rgb = RGBColor(30, 58, 138)

    table_sec = doc.add_table(rows=11, cols=4)
    table_sec.alignment = WD_TABLE_ALIGNMENT.CENTER
    sec_headers = ["STT", "Đối Tượng / Hành Vi Thử Nghiệm", "Kỳ Vọng Phân Quyền", "Kết Quả Thực Tế"]
    for i, h in enumerate(sec_headers):
        cell = table_sec.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    sec_data = [
        ("1", "Unauthenticated User", "Cho phép xem bài học miễn phí", "🛡️ PASS (200 OK)"),
        ("2", "Enrolled Student", "Cho phép hỏi đáp bài học đã đăng ký", "🛡️ PASS (200 OK + Verified Sources)"),
        ("3", "Non-enrolled Student", "Chặn hỏi đáp bài học trả phí", "🛡️ PASS (403 Forbidden)"),
        ("4", "Admin", "Toàn quyền truy cập tất cả bài học", "🛡️ PASS (Full Access)"),
        ("5", "Course Instructor", "Toàn quyền trên khóa học mình phụ trách", "🛡️ PASS (Full Access)"),
        ("6", "Tampered Lesson ID (999999)", "Chặn truy cập ID không tồn tại", "🛡️ PASS (404 Not Found)"),
        ("7", "Tampered Course ID", "Chặn liên kết sai lệch khóa học", "🛡️ PASS (Ràng buộc CSDL thật)"),
        ("8", "Tampered Source Card Injection", "Loại bỏ card giả mạo không có trong DB", "🛡️ PASS (0 fake cards)"),
        ("9", "Cross-User History Attempt", "User 1 không đọc được chat User 2", "🛡️ PASS (100% History Isolated)"),
        ("10", "Cross-Course Retrieval Attempt", "Không tìm thấy bài học Course khác", "🛡️ PASS (100% Course Boundary)")
    ]

    for row_idx, data in enumerate(sec_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_sec.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 60, 60, 80, 80)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx in [0, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                if col_idx == 3:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_paragraph()

    # 6. HIỆU NĂNG & HỒ SƠ GỌI API
    h6 = doc.add_heading(level=1)
    r_h6 = h6.add_run("6. ĐO LƯỜNG HIỆU NĂNG THỰC TẾ & HỒ SƠ GỌI API")
    r_h6.font.name = "Arial"
    r_h6.font.color.rgb = RGBColor(30, 58, 138)

    p_perf = doc.add_paragraph()
    r = p_perf.add_run(
        "Đo đạc trên tập mẫu 12 lượt Streaming Requests thực tế bao phủ 4 nhóm câu hỏi (Current Lesson QA, Course-Wide Search, Follow-up Rewrite, Timestamp QA):\n"
        "• Time To First Token (TTFT): Avg = 2.18s | P50 = 1.97s | P95 = 2.79s\n"
        "• Total Completion Latency: Avg = 3.46s | P50 = 3.12s | P95 = 4.84s\n"
        "• Pure Timestamp Added Latency: Avg = 234.57 ms (< 300ms, truy vấn Index PostgreSQL siêu tốc).\n"
        "• Tỉ lệ Fast-Path Intent: ~85.0% (Phản hồi < 1ms, không tiêu tốn token) | LLM Fallback Intent: ~15.0%.\n"
        "• Thống kê gọi API trung bình:\n"
        "  - Current Lesson QA: 0 Embeddings | 1 Gemini Gen | 0 Fallback | 0 Pinecone | 2 PostgreSQL\n"
        "  - Course-Wide Search: 1 Embedding | 1 Gemini Gen | 0-1 Fallback | 0-1 Pinecone | 2-3 PostgreSQL\n"
        "  - Follow-up Rewrite: 1 Embedding | 1 Gemini Gen | 1 Rewriter | 0-1 Pinecone | 3 PostgreSQL\n"
        "  - Timestamp Video QA: 0 Embeddings | 1 Gemini Gen | 0 Fallback | 0 Pinecone | 2 PostgreSQL"
    )
    r.font.name = "Arial"
    r.font.size = Pt(9.5)

    doc.add_paragraph()

    # 7. QUAN SÁT & BẢO MẬT NHẬT KÝ
    h7 = doc.add_heading(level=1)
    r_h7 = h7.add_run("7. NHẬT KÝ QUAN SÁT AN TOÀN & CHẤT LƯỢNG BUILD")
    r_h7.font.name = "Arial"
    r_h7.font.color.rgb = RGBColor(30, 58, 138)

    p_obs = doc.add_paragraph()
    r = p_obs.add_run(
        "• Observability: Development logs ghi nhận đầy đủ Request ID, Intent, Scope, Rewritten Flag, Source Count, Latency Breakdown.\n"
        "• Safe Logging: Tuyệt đối KHÔNG ghi nhật ký JWT Token, API Key, Mật khẩu hay thông tin nhạy cảm của người dùng.\n"
        "• Frontend Build (npm run build): PASS (✓ built in 12.18s — 0 lỗi, 0 cảnh báo syntax).\n"
        "• Backend Startup: PASS (Khởi động sạch sẽ, 0 circular dependency).\n"
        "• Linter Tool: NOT CONFIGURED (Sử dụng tiêu chuẩn mặc định của Vite/Babel)."
    )
    r.font.name = "Arial"
    r.font.size = Pt(9.5)

    doc.add_paragraph()

    # 8. GIỚI HẠN KỸ THUẬT & KẾT LUẬN
    h8 = doc.add_heading(level=1)
    r_h8 = h8.add_run("8. GIỚI HẠN KỸ THUẬT & KẾT LUẬN NGHIỆM THU")
    r_h8.font.name = "Arial"
    r_h8.font.color.rgb = RGBColor(30, 58, 138)

    p_concl = doc.add_paragraph()
    r = p_concl.add_run(
        "• Known Limitations: Quy mô dataset hiện tại (~50 bài học, 100+ subtitle cues); cơ chế multi-tab chia sẻ chung rolling 30-min window.\n"
        "• P0 Blockers: 0 | P1 Blockers: 0 | P2 Backlog Items: 2 (Bổ sung UUID conversation_id cho từng tab; Redis Semantic Cache).\n\n"
        "KẾT LUẬN: Hệ thống đạt trạng thái PASS trên toàn bộ Acceptance/Evaluation Suite và ĐẠT CHUẨN PRODUCTION-READINESS trong phạm vi dữ liệu và yêu cầu của đồ án tốt nghiệp."
    )
    r.font.name = "Arial"
    r.font.size = Pt(9.5)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s = p_sign.add_run("\nTP. Hồ Chí Minh, ngày 18 tháng 08 năm 2026\nĐẠI DIỆN NHÓM PHÁT TRIỂN DỰ ÁN\n\n\n")
    r_s.font.name = "Arial"
    r_s.font.size = Pt(10)
    r_s.font.italic = True
    
    r_signers = p_sign.add_run("NGUYỄN DŨNG QUỐC ANH — NGUYỄN THANH LIÊM — LÊ ĐÌNH CHƯƠNG")
    r_signers.font.name = "Arial"
    r_signers.font.bold = True
    r_signers.font.size = Pt(10)
    r_signers.font.color.rgb = RGBColor(30, 58, 138)

    output_path = "e:/Project-E-learning-website-for-learning-English-online/BAO_CAO_PHASE9_FINAL_PRODUCTION_CLOSEOUT.docx"
    doc.save(output_path)
    print("Report generated successfully at: " + output_path)

if __name__ == "__main__":
    create_report()
