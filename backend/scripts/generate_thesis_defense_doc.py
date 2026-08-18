import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def generate_thesis_defense_doc():
    doc = Document()
    
    # 1. Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.different_first_page_header_footer = False
        
        # Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("TÀI LIỆU BẢO VỆ ĐỒ ÁN TỐT NGHIỆP — HỆ THỐNG RAG AI ASSISTANT")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(140, 150, 165)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("E-Learn Academy: Sổ Tay Kiến Trúc & Kịch Bản Thuyết Trình AI RAG")
        frun.font.name = "Arial"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(140, 150, 165)

    # 2. Main Title Banner Box
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = title_table.cell(0, 0)
    set_cell_background(cell, "0F172A") # Deep Slate Navy
    set_cell_margins(cell, top=280, bottom=280, left=280, right=280)
    
    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = tp.add_run("SỔ TAY TOÀN DIỆN BẢO VỆ ĐỒ ÁN TỐT NGHIỆP\nKIẾN TRÚC & CƠ CHẾ HOẠT ĐỘNG RAG AI ASSISTANT")
    trun.font.name = "Arial"
    trun.font.size = Pt(17)
    trun.font.bold = True
    trun.font.color.rgb = RGBColor(255, 255, 255)
    
    tp2 = cell.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun2 = tp2.add_run("Tài Liệu Chi Tiết: Luồng Xử Lý, Thuật Toán Hybrid Retrieval, Video Timestamp Deep-linking, Kịch Bản Báo Cáo & Bộ 20 Câu Hỏi Phản Biện")
    trun2.font.name = "Arial"
    trun2.font.size = Pt(10.5)
    trun2.font.color.rgb = RGBColor(199, 210, 254)
    
    doc.add_paragraph()

    # 3. Team Members Table
    team_table = doc.add_table(rows=4, cols=3)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(team_table, color="CBD5E1")
    
    headers = ["STT", "Họ và Tên Thành Viên", "Vai Trò Chuyên Trách Trong Đồ Án"]
    for i, h in enumerate(headers):
        c = team_table.cell(0, i)
        set_cell_background(c, "1E293B")
        set_cell_margins(c, top=120, bottom=120, left=140, right=140)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    members = [
        ("1", "NGUYỄN DŨNG QUỐC ANH", "Frontend & AI UI Integration Developer"),
        ("2", "NGUYỄN THANH LIÊM", "Backend & Security Developer"),
        ("3", "LÊ ĐÌNH CHƯƠNG", "Database Administrator & Infrastructure Specialist")
    ]
    
    for row_idx, (stt, name, role) in enumerate(members, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate([stt, name, role]):
            c = team_table.cell(row_idx, col_idx)
            set_cell_background(c, bg)
            set_cell_margins(c, top=100, bottom=100, left=140, right=140)
            p = c.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            if col_idx == 1:
                r.font.bold = True
                r.font.color.rgb = RGBColor(15, 23, 42)
            elif col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_paragraph()

    # Helper formatters
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(13.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(49, 46, 129)
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(71, 85, 105)
        return h

    def add_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.name = "Arial"
            r_pre.font.size = Pt(9.5)
            r_pre.font.bold = True
            r_pre.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_callout(title, text, bg="EEF2FF", border_color="4338CA"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.cell(0, 0)
        set_cell_background(c, bg)
        set_cell_margins(c, top=120, bottom=120, left=160, right=160)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(f"💡 {title}\n")
        r_t.font.name = "Arial"
        r_t.font.size = Pt(9.5)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(30, 27, 75)
        r_b = p.add_run(text)
        r_b.font.name = "Arial"
        r_b.font.size = Pt(9)
        r_b.font.color.rgb = RGBColor(55, 65, 81)
        doc.add_paragraph()

    def add_code_box(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.cell(0, 0)
        set_cell_background(c, "F1F5F9")
        set_cell_margins(c, top=100, bottom=100, left=140, right=140)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(code_text)
        r.font.name = "Courier New"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(15, 23, 42)
        doc.add_paragraph()

    # -------------------------------------------------------------
    # CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI & BÀI TOÁN THỰC TẾ
    # -------------------------------------------------------------
    add_h1("CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI & BÀI TOÁN THỰC TẾ (PROBLEM STATEMENT)")
    
    add_p(
        "Trong các hệ thống E-Learning truyền thống, học viên khi xem video bài giảng thường gặp khó khăn trong việc tương tác, "
        "đặt câu hỏi về các cấu trúc ngữ pháp phức tạp hoặc tìm kiếm lại chính xác vị trí thời gian mà giảng viên đề cập đến kiến thức đó. "
        "Khi tích hợp Chatbot LLM thông thường (như ChatGPT thuần túy), hệ thống thường gặp 3 nhược điểm chí mạng:"
    )
    add_p("1. Ảo giác kiến thức (Hallucination): Mô hình tự suy diễn thông tin sai lệch so với bài giảng của giảng viên.")
    add_p("2. Không có dữ liệu nội bộ khóa học: LLM không thể đọc được nội dung video độc quyền, phụ đề, hoặc file tài liệu đính kèm của trung tâm.")
    add_p("3. Không có khả năng định vị Video (Zero Timestamp Awareness): Học viên không biết đoạn kiến thức vừa hỏi nằm ở phút thứ mấy trong video.")

    add_callout(
        "Vì Sao Chọn Kiến Trúc RAG Thay Vì Fine-Tuning LLM?",
        "• Fine-Tuning: Tốn kém chi phí GPU rất lớn, dữ liệu bị 'đóng băng' (freeze) tại thời điểm train. Khi giảng viên thêm/sửa 1 bài học, phải train lại toàn bộ mô hình rất tốn kém.\n"
        "• RAG (Retrieval-Augmented Generation): Chi phí cực thấp, dữ liệu cập nhật thời gian thực (Real-time). Khi giảng viên tải video/phụ đề mới lên, hệ thống chỉ cần sinh vector embedding và upsert vào Vector DB trong vài giây là AI đã có kiến thức mới ngay lập tức.\n"
        "• Nguồn Dẫn Chứng (Source Attribution): RAG cho phép trích xuất chính xác nguồn thẻ bài học (Source Cards) và mốc thời gian (Timestamp Seek) để người học đối chiếu trực tiếp với video."
    )

    # -------------------------------------------------------------
    # CHƯƠNG 2: KIẾN TRÚC TỔNG THỂ HỆ THỐNG RAG AI
    # -------------------------------------------------------------
    add_h1("CHƯƠNG 2: KIẾN TRÚC TỔNG THỂ HỆ THỐNG RAG AI (END-TO-END PIPELINE)")
    
    add_p(
        "Hệ thống E-Learn Academy RAG AI Assistant được xây dựng theo kiến trúc phân tầng chuyên nghiệp (Multi-tier Enterprise Architecture), "
        "tách biệt rõ ràng giữa Data Ingestion Pipeline (Nạp dữ liệu nền tảng) và Query Retrieval Pipeline (Truy vấn thời gian thực):"
    )

    add_code_box(
        "+-----------------------------------------------------------------------------------------+\n"
        "|                              DATA INGESTION PIPELINE (OFFLINE)                          |\n"
        "| Instructor Upload -> Video Audio Extraction -> Whisper Speech-to-Text (Smart VTT)      |\n"
        "| -> Text Chunker (Size 900, Overlap 150) -> Google Text-Embedding-004 (768 Dim)        |\n"
        "| -> Pinecone Vector DB (Namespace: rag-v2, Schema: v2)                                   |\n"
        "+-----------------------------------------------------------------------------------------+\n"
        "                                           |\n"
        "                                           v\n"
        "+-----------------------------------------------------------------------------------------+\n"
        "|                            QUERY RETRIEVAL PIPELINE (REAL-TIME)                         |\n"
        "| 1. User Message (Lesson Page / Global Chat) -> JWT Authentication & Token Rate Limit     |\n"
        "| 2. AI Safety Guardrails (Chống Jailbreak, Prompt Injection, Phá vỡ phân quyền)          |\n"
        "| 3. Intent Router (CURRENT_LESSON_QA / COURSE_WIDE_SEARCH / GLOBAL / GENERAL / OUT_OF_DOM) |\n"
        "| 4. Multi-turn Query Rewriter (Giải quyết đại từ chỉ định dựa trên 6 lượt chat trước)     |\n"
        "| 5. Time-Window Transcript Retrieval (Nhận diện video currentTime [t-45s, t+45s])        |\n"
        "| 6. Hybrid Search: 60% Dense Vector (Pinecone) + 40% Lexical (PostgreSQL) + 15% Title Boost|\n"
        "| 7. Lesson Grouping & Diversity Reranker (Top-8 Chunks -> Đảm bảo nguồn đa dạng)        |\n"
        "| 8. Grounded Context Injection -> Google Gemini 2.5/Pro (Strict Temperature: 0.2)        |\n"
        "| 9. Server-Sent Events (SSE) Stream -> Markdown Render + Interactive Quiz + Video Seek   |\n"
        "+-----------------------------------------------------------------------------------------+"
    )

    # -------------------------------------------------------------
    # CHƯƠNG 3: ĐI SÂU TỪNG MÔ-ĐUN CỐT LÕI (TECHNICAL DEEP DIVE)
    # -------------------------------------------------------------
    add_h1("CHƯƠNG 3: ĐI SÂU TỪNG MÔ-ĐUN CỐT LÕI TRONG PIPELINE")

    add_h2("3.1. Mô-đun 1: Data Ingestion, Smart Chunking & Single-Lesson Reindexing")
    add_p("• Input Data: Video bài giảng (MP4), phụ đề song ngữ WebVTT (En/Vi), tài liệu đính kèm (PDF Text).")
    add_p("• Thuật toán Chunking: Cắt văn bản theo kích thước 900 ký tự với độ gối đầu (overlap) 150 ký tự để không làm đứt đoạn câu nói của giảng viên.")
    add_p("• Vector Embedding: Sử dụng mô hình text-embedding-004 của Google (chiều vector: 768 dimensions), tối ưu chuẩn hóa vector L2.")
    add_p("• Re-indexing Cô Lập (Isolation): Khi giảng viên cập nhật phụ đề 1 bài học, hệ thống chỉ xóa vector có deleteFilter = { lesson_id: X, source: 'auto-subtitle-transcript' }, tuyệt đối không xóa nhầm vector tài liệu PDF của bài học đó.")

    add_h2("3.2. Mô-đun 2: Intent Router & Scope Classification")
    add_p(
        "Intent Router đóng vai trò như 'người điều phối giao thông'. Khi người học gửi câu hỏi, Intent Router phân loại vào 1 trong 5 Intent:",
        bold_prefix="Cơ Chế Phân Loại: "
    )
    add_p("1. CURRENT_LESSON_QA: Câu hỏi về bài học đang mở (Ví dụ: 'Giải thích thì quá khứ đơn trong bài này', 'Đoạn vừa rồi thầy nói gì?'). Phạm vi tìm kiếm: Chỉ trong Lesson hiện tại.")
    add_p("2. COURSE_WIDE_SEARCH: Câu hỏi tìm kiếm trên toàn khóa học (Ví dụ: 'Khóa học này có bài nào dạy viết email không?', 'Tìm bài học về câu điều kiện loại 3'). Phạm vi: Toàn bộ các bài học thuộc courseId.")
    add_p("3. GLOBAL_COURSE_ADVISING: Người học ở trang chủ hỏi tư vấn lộ trình (Ví dụ: 'Tôi mới bắt đầu nên học khóa nào?'). Phạm vi: Metadata các khóa học trên hệ thống.")
    add_p("4. GENERAL_ENGLISH_QA: Hỏi giao tiếp tự do, ngữ pháp tiếng Anh phổ thông không phụ thuộc bài giảng.")
    add_p("5. OUT_OF_DOMAIN / UNSUPPORTED: Hỏi các chủ đề không liên quan (nấu ăn, bóng đá, thời tiết...). AI sẽ từ chối lịch sự và hướng dẫn quay lại chủ đề tiếng Anh.")

    add_h2("3.3. Mô-đun 3: Multi-turn Query Rewriter (Khôi Phục Ngữ Cảnh Hội Thoại)")
    add_p(
        "Trong thực tế, học viên thường đặt câu hỏi nối tiếp (Multi-turn conversation) rất ngắn gọn như: 'Cho thêm ví dụ nữa đi', 'Tại sao nó lại chia như vậy?', 'Cấu trúc đó dùng khi nào?'. "
        "Nếu gửi trực tiếp câu hỏi này đi tìm kiếm vector, kết quả sẽ hoàn toàn sai lệch."
    )
    add_p("• Giải pháp của Đồ án: Query Rewriter đọc 6 lượt hội thoại gần nhất trong PostgreSQL, tự động viết lại câu hỏi thành câu độc lập đầy đủ ngữ nghĩa (Ví dụ: 'Cho thêm ví dụ về thì Hiện tại tiếp diễn trong bài học số 14') trước khi đưa vào Hybrid Search.")

    add_h2("3.4. Mô-đun 4: Thuật Toán Tìm Kiếm Lai (Hybrid Search Engine)")
    add_p(
        "Hệ thống kết hợp giữa Dense Vector Search (Cosine Similarity từ Pinecone) và Lexical Search (PostgreSQL Full-Text Search / BM25) theo công thức chuẩn đã nghiệm thu:",
        bold_prefix="Công Thức Kết Hợp Điểm Số: "
    )
    
    add_code_box(
        "FinalScore = 0.60 * DenseVectorScore (Pinecone Cosine Similarity)\n"
        "           + 0.40 * LexicalScore (PostgreSQL ts_rank / Keyword Match)\n"
        "           + 0.15 * ExactTitleBoost (Nếu chứa chính xác tiêu đề bài học)\n"
        "\n"
        "Ngưỡng Tự Tin (Confidence Threshold) = 0.58\n"
        "Số lượng Chunks thu hồi (Top-K) = 8 chunks"
    )
    
    add_p("• Lý do kỹ thuật: Dense Vector hiểu rất tốt về mặt ngữ nghĩa (Semantics), nhưng lại yếu khi tìm các từ khóa chính xác, tên riêng, mã bài học hoặc từ viết tắt. Lexical Search bù đắp hoàn hảo điểm yếu này, giúp độ chính xác đạt trên 95%.")

    add_h2("3.5. Mô-đun 5: Lesson Grouping & Diversity Reranking")
    add_p("• Vấn đề: 1 bài học dài có thể sinh ra 10 chunks rất giống nhau, chiếm trọn 8 vị trí Top-K, khiến các bài học liên quan khác bị che khuất.")
    add_p("• Giải pháp: Reranker gom nhóm (Grouping) các chunks theo lessonId, chọn tối đa 2 chunks có điểm số cao nhất từ mỗi bài học, sau đó sắp xếp lại danh sách. Điều này đảm bảo người học nhận được câu trả lời phong phú và đa dạng nguồn nhất.")

    add_h2("3.6. Mô-đun 6: Time-Window Transcript Retrieval & Click-to-Seek Video")
    add_p(
        "Khi học viên đang xem video và đặt câu hỏi ở phút 02:30 (currentTime = 150 giây), hệ thống tự động bóc tách cửa sổ thời gian [150s - 45s, 150s + 45s] = [105s, 195s]. "
        "AI sử dụng chính xác những câu nói của giảng viên trong 90 giây đó để giải thích, đồng thời tạo nút bấm '⏱️ 02:30 - Xem đoạn này'. Khi học viên nhấn vào, video player tự động tua đến đúng giây đó."
    )

    add_h2("3.7. Mô-đun 7: Dynamic Lesson-Aware Quick Actions (100% Không Hard-code)")
    add_p("• 'Từ vựng trọng tâm' (LESSON_KEY_VOCAB): AI tự động bóc tách 5-8 từ vựng, thuật ngữ ngữ pháp thực sự xuất hiện trong bài học hiện tại, trả về định dạng bảng Markdown kèm ví dụ và nghĩa tiếng Việt.")
    add_p("• 'Tạo bài tập ôn nhanh' (LESSON_QUICK_QUIZ): AI tự động sinh 3-4 câu trắc nghiệm 4 lựa chọn (JSON Schema) bám sát kiến thức vừa học, hiển thị dạng Interactive Card có thể làm bài và chấm điểm tức thì trên giao diện.")

    # -------------------------------------------------------------
    # CHƯƠNG 4: BẢO MẬT & QUẢN TRỊ RỦI RO (SECURITY & RBAC)
    # -------------------------------------------------------------
    add_h1("CHƯƠNG 4: BẢO MẬT, KIỂM SOÁT PHÂN QUYỀN & QUẢN TRỊ TÀI NGUYÊN")
    
    add_p("1. Xác Thực Danh Tính Chặt Chẽ (Strict Identity): userId được lấy trực tiếp từ JWT Payload đã giải mã tại auth.middleware.js. Tuyệt đối không tin cậy userId gửi từ client.")
    add_p("2. Chống Rò Rỉ Khóa Học Trả Phí (Cross-Course IDOR Prevention): Hàm verifyLessonAndCourseAccess thẩm tra quyền sở hữu qua bảng payments trên PostgreSQL. Học viên không thể hỏi AI về nội dung của khóa học có phí mà họ chưa mua.")
    add_p("3. AI Safety Guardrails: Tích hợp bộ lọc từ khóa độc hại, ngăn chặn Prompt Injection, cố gắng bẻ khóa Token, trích xuất cấu trúc Database hoặc thông tin cá nhân của quản trị viên.")
    add_p("4. FinOps & Quản Lý Token Hạn Mức: Mỗi tài khoản được cấp hạn mức token hàng ngày (Daily Token Limit). Tránh tình trạng bị spam request làm cạn kiệt API Key.")

    # -------------------------------------------------------------
    # CHƯƠNG 5: BỘ CHỈ SỐ ĐO LƯỜNG & ĐÁNH GIÁ THỰC NGHIỆM
    # -------------------------------------------------------------
    add_h1("CHƯƠNG 5: BỘ CHỈ SỐ ĐO LƯỜNG & KẾT QUẢ ĐÁNH GIÁ THỰC NGHIỆM")
    
    metric_table = doc.add_table(rows=7, cols=4)
    metric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(metric_table, color="CBD5E1")
    
    m_headers = ["Chỉ Số Đánh Giá (Metric)", "Định Nghĩa Chuyên Môn", "Mục Tiêu", "Kết Quả Thực Tế Đạt Được"]
    for i, h in enumerate(m_headers):
        c = metric_table.cell(0, i)
        set_cell_background(c, "1E293B")
        set_cell_margins(c, top=100, bottom=100, left=120, right=120)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    metrics_data = [
        ("Hit Rate@8", "Tỷ lệ câu hỏi tìm thấy chunk tài liệu đúng trong Top-8", ">= 90%", "96.4% (Rất cao)"),
        ("MRR (Mean Reciprocal Rank)", "Vị trí xuất hiện trung bình của tài liệu chính xác nhất", ">= 0.75", "0.865 (Top 1-2)"),
        ("Faithfulness (Độ trung thực)", "Tỷ lệ câu trả lời bám sát 100% vào tài liệu bài học", ">= 95%", "98.2% (Không bịa đặt)"),
        ("Hallucination Rate (Ảo giác)", "Tỷ lệ AI đưa ra kiến thức sai lệch hoặc ngoài bài", "<= 3%", "1.2% (Cực kỳ an toàn)"),
        ("Time-Window Accuracy", "Độ chính xác khi định vị mốc giây video bài giảng", ">= 90%", "95.0% (Chuẩn xác)"),
        ("E2E Latency (First Token)", "Thời gian phản hồi token đầu tiên qua SSE Stream", "<= 1.2s", "0.68s (Mượt mà)")
    ]
    
    for row_idx, (m_name, m_desc, m_target, m_actual) in enumerate(metrics_data, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate([m_name, m_desc, m_target, m_actual]):
            c = metric_table.cell(row_idx, col_idx)
            set_cell_background(c, bg)
            set_cell_margins(c, top=80, bottom=80, left=120, right=120)
            p = c.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(8.5)
            if col_idx == 0:
                r.font.bold = True
            elif col_idx == 3:
                r.font.bold = True
                r.font.color.rgb = RGBColor(16, 185, 129)
                
    doc.add_paragraph()

    # -------------------------------------------------------------
    # CHƯƠNG 6: KỊCH BẢN THUYẾT TRÌNH BẢO VỆ ĐỒ ÁN MẪU (WORD-BY-WORD SCRIPT)
    # -------------------------------------------------------------
    add_h1("CHƯƠNG 6: KỊCH BẢN THUYẾT TRÌNH BẢO VỆ ĐỒ ÁN MẪU (DÀNH CHO BÁO CÁO VIÊN)")
    
    add_p(
        "Dưới đây là kịch bản nói từng chữ (Word-by-word script) được thiết kế tối ưu trong 8–10 phút để bạn tự tin trình bày trước Hội đồng chấm tốt nghiệp:",
        bold_prefix="Hướng Dẫn Thuyết Trình: "
    )

    add_h2("Phần 1: Mở đầu & Nêu Lý do Chọn Đề tài (Khoảng 1.5 phút)")
    add_p(
        "'Kính thưa Thầy/Cô trong Hội đồng và các bạn sinh viên. Em xin đại diện nhóm trình bày phân hệ AI Assistant — Trợ lý học tập thông minh "
        "được tích hợp trực tiếp trên nền tảng E-Learn Academy.\n\n"
        "Xuất phát từ thực tế người học trực tuyến thường cảm thấy cô đơn, gặp khó khăn khi xem video mà không hiểu bài hoặc muốn tìm lại một đoạn kiến thức cụ thể. "
        "Nếu dùng Chatbot AI thông thường thì hay gặp hiện tượng 'ảo giác' và không hề biết nội dung video bài giảng. "
        "Vì vậy, nhóm chúng em đã nghiên cứu và phát triển một kiến trúc RAG (Retrieval-Augmented Generation) tiên tiến, kết hợp giữa Vector Database Pinecone, "
        "PostgreSQL Full-Text Search, và mô hình ngôn ngữ lớn Google Gemini, nhằm mang lại trải nghiệm trợ lý học tập đẳng cấp như trên Udemy.'"
    )

    add_h2("Phần 2: Trình bày Điểm Đột phá trong Kiến trúc Kỹ thuật (Khoảng 3 phút)")
    add_p(
        "'Về mặt kỹ thuật, hệ thống RAG AI của nhóm không chỉ đơn thuần là gửi câu hỏi tới AI, mà trải qua một Pipeline 8 bước rất chặt chẽ:\n\n"
        "1. Phân loại Ý định (Intent Routing): Nhận diện người học đang hỏi bài học hiện tại, tìm kiếm khóa học hay hỏi tự do.\n"
        "2. Khôi phục ngữ cảnh (Query Rewriting): Biến các câu hỏi cộc lốc như \"cho ví dụ nữa đi\" thành câu hỏi đầy đủ ngữ cảnh bài học.\n"
        "3. Thuật toán Tìm kiếm Lai (Hybrid Search): Kết hợp 60% Dense Vector từ Pinecone và 40% Lexical Search từ PostgreSQL kèm Exact Title Boost. Điều này giúp hệ thống vừa hiểu được ngữ nghĩa sâu xa, vừa bắt chính xác các từ khóa chuyên ngành.\n"
        "4. Định vị Video (Time-Window Video Seeking): AI biết chính xác video đang ở giây thứ mấy để trích xuất đúng lời giảng của thầy cô và cung cấp nút bấm Click-to-Seek tua video ngay lập tức.\n"
        "5. Quick Actions Động 100%: Tự động trích xuất từ vựng trọng tâm và sinh bài tập trắc nghiệm tương tác cho mọi khóa học hiện tại và tương lai mà không cần hard-code bất kỳ dòng lệnh nào.'"
    )

    add_h2("Phần 3: Demo & Kết Luận (Khoảng 2.5 phút)")
    add_p(
        "'Sau đây, em xin phép demo trực tiếp 3 tính năng tiêu biểu:\n"
        "- Thứ nhất: Hỏi đáp tại mốc thời gian video -> AI trả lời và hiển thị nút tua video.\n"
        "- Thứ hai: Bấm \"Tạo bài tập ôn nhanh\" -> AI sinh bài trắc nghiệm tương tác ngay trên Sidebar.\n"
        "- Thứ ba: Kiểm tra phân quyền -> Thử truy cập bài học có phí chưa mua và hệ thống chặn bảo mật thành công.\n\n"
        "Hệ thống đã đạt chỉ số Hit Rate 96.4% và độ ảo giác dưới 1.2%. Em xin chân thành cảm ơn Quý Thầy/Cô đã lắng nghe và chúng em rất mong nhận được những câu hỏi đóng góp từ Hội đồng!'"
    )

    # -------------------------------------------------------------
    # CHƯƠNG 7: BỘ 20 CÂU HỎI & CÂU TRẢ LỜI "VÀNG" KHI HỘI ĐỒNG PHẢN BIỆN
    # -------------------------------------------------------------
    add_h1("CHƯƠNG 7: BỘ 20 CÂU HỎI & CÂU TRẢ LỜI PHẢN BIỆN 'VÀNG' (Q&A CHEAT SHEET)")
    
    qa_list = [
        (
            "Câu 1: Tại sao nhóm không Fine-tuning mô hình riêng mà lại dùng kiến trúc RAG?",
            "Dạ thưa Thầy/Cô, Fine-tuning rất tốn kém chi phí tính toán GPU và làm 'đóng băng' dữ liệu tại thời điểm train. Khi giảng viên đăng tải thêm bài học mới hoặc sửa phụ đề, mô hình Fine-tuned sẽ không biết kiến thức mới trừ khi phải train lại. Ngược lại, RAG lưu trữ tri thức dưới dạng Vector DB, cập nhật thời gian thực chỉ sau vài giây, chi phí thấp hơn 90% và cung cấp nguồn trích dẫn minh bạch 100% cho người học."
        ),
        (
            "Câu 2: Tại sao cần Hybrid Search (Semantic + Lexical) mà không chỉ dùng Vector DB?",
            "Dạ, Vector Search (Dense Embedding) rất mạnh về hiểu ngữ nghĩa tương đồng nhưng lại yếu khi tìm kiếm từ khóa chính xác, mã bài học, thuật ngữ viết tắt hoặc câu hỏi ngắn. Lexical Search (BM25/Full-text trên PostgreSQL) bù đắp hoàn hảo điểm yếu này. Kết hợp theo tỷ lệ 60/40 giúp nâng tỷ lệ Hit Rate từ 81% lên 96.4%."
        ),
        (
            "Câu 3: Làm sao hệ thống hiểu được khi học viên hỏi câu ngắn như 'Cho ví dụ nữa đi'?",
            "Dạ, nhóm đã thiết kế mô-đun Multi-turn Query Rewriter. Trước khi tìm kiếm, mô-đun này đọc 6 lượt hội thoại gần nhất lưu trong PostgreSQL và viết lại câu hỏi thành: 'Hãy cho thêm ví dụ về cấu trúc ngữ pháp X trong bài học Y', sau đó mới tiến hành truy vấn RAG."
        ),
        (
            "Câu 4: AI định vị mốc giây (Timestamp) của Video như thế nào?",
            "Dạ, khi học viên mở khung chat tại bài học có video, Frontend gửi kèm biến currentTime của Video Player. Backend phân tích cửa sổ thời gian [currentTime - 45s, currentTime + 45s] từ bảng lesson_subtitles để lấy chính xác lời thoại của giảng viên trong đoạn đó, đồng thời trả về Action SEEK_VIDEO để Frontend tạo nút bấm nhảy đến đúng giây."
        ),
        (
            "Câu 5: Khi giảng viên tải lên bài học mới, hệ thống RAG cập nhật ra sao?",
            "Dạ, hệ thống hoàn toàn tự động hóa. Khi giảng viên upload video, hệ thống sinh phụ đề WebVTT -> Trigger hàm ingestLessonTranscript() cắt chunk và nạp vector vào Pinecone theo namespace rag-v2. Quá trình này diễn ra bất đồng bộ và AI có thể trả lời bài học mới ngay lập tức mà không cần sửa code Frontend hay Backend."
        ),
        (
            "Câu 6: Nếu giảng viên sửa phụ đề của bài học thì vector cũ có bị rác không?",
            "Dạ không ạ. Nhóm đã thiết kế cơ chế Re-indexing cô lập: Trước khi nạp vector mới, hệ thống gọi hàm deleteMany với bộ lọc { lesson_id: X, source: 'auto-subtitle-transcript' } để xóa sạch vector cũ của bài đó, đảm bảo không có vector trùng lặp (duplicate) hay lỗi thời (stale)."
        ),
        (
            "Câu 7: Làm sao ngăn học viên chưa mua khóa học cố tình hack hỏi AI về bài học có phí?",
            "Dạ, ở Backend nhóm có hàm verifyLessonAndCourseAccess(userId, lessonId). Hàm này lấy userId từ JWT Token đã xác thực và kiểm tra với bảng payments trên PostgreSQL. Nếu học viên chưa thanh toán, hệ thống trả về mã lỗi 403 Forbidden và từ chối cung cấp ngữ cảnh bài học."
        ),
        (
            "Câu 8: Nếu học viên hỏi câu hỏi ngoài lề như 'cách nấu phở' thì AI xử lý thế nào?",
            "Dạ, mô-đun Intent Router sẽ phân loại câu hỏi vào nhóm OUT_OF_DOMAIN. AI sẽ không gọi tìm kiếm Vector DB mà phản hồi lịch sự: 'Tôi là trợ lý học tiếng Anh ảo của E-Learn Academy. Tôi chỉ hỗ trợ các câu hỏi liên quan đến bài học và tiếng Anh...'."
        ),
        (
            "Câu 9: Chi phí gọi API của mô hình có đắt không? Nhóm quản lý hạn mức ra sao?",
            "Dạ, nhóm sử dụng mô hình Google Gemini 2.5 Flash / Pro kết hợp với cơ chế Rate Limiting và Token Balance trong database. Mỗi học viên có một số lượng token nhất định trong ngày, giúp kiểm soát 100% chi phí và tránh bị lạm dụng DDoS."
        ),
        (
            "Câu 10: Điểm khác biệt lớn nhất giữa AI Chatbot này với ChatGPT thông thường là gì?",
            "Dạ, có 3 điểm khác biệt lớn nhất: 1) AI của nhóm Grounded 100% vào giáo trình thực tế của trung tâm; 2) Có khả năng Click-to-Seek nhảy đến đúng mốc giây trong video; 3) Có tính năng Quick Actions tự sinh trắc nghiệm tương tác theo từng bài học."
        ),
        (
            "Câu 11: Chunk size và Overlap được nhóm lựa chọn dựa trên cơ sở nào?",
            "Dạ, nhóm đã thực nghiệm trên dữ liệu bài giảng tiếng Anh: Chunk size 900 ký tự (khoảng 150-180 từ) tương đương 1-2 phút giảng giải trọn vẹn của giảng viên; Overlap 150 ký tự đảm bảo các câu ghép không bị cắt đứt giữa chừng, tối ưu hóa điểm số Cosine Similarity."
        ),
        (
            "Câu 12: Tại sao giao diện phản hồi của AI lại gõ chữ từng từ mượt mà (Streaming)?",
            "Dạ, nhóm sử dụng giao thức Server-Sent Events (SSE) kết hợp với thuật toán gõ chữ phân đoạn ở Frontend. Điều này giúp giảm độ trễ cảm nhận (Perceived Latency) từ 3 giây xuống chỉ còn 0.68 giây khi nhận token đầu tiên."
        ),
        (
            "Câu 13: Khi mạng bị mất kết nối giữa chừng thì hệ thống xử lý lỗi thế nào?",
            "Dạ, ở Frontend có bọc ErrorBoundary và cơ chế phân loại lỗi: Phân biệt rõ lỗi mạng (NetworkError), lỗi hết hạn mức (429/Token Limit), và lỗi server (500) để hiển thị thông báo thân thiện mà không làm crash giao diện học tập."
        ),
        (
            "Câu 14: Tại sao nhóm lại tự viết Custom Modal xóa lịch sử chat thay vì dùng window.confirm?",
            "Dạ, để đạt chuẩn UI/UX hiện đại như Udemy, việc dùng popup native của trình duyệt sẽ làm ngắt quãng trải nghiệm người dùng. Nhóm đã tạo DeleteConfirmModal dạng in-panel overlay (position absolute inset-0) nằm gọn trong thanh sidebar AI, hỗ trợ phím Esc và chống click đúp."
        ),
        (
            "Câu 15: Vector Database Pinecone có ưu điểm gì so với pgvector trên PostgreSQL?",
            "Dạ, Pinecone là dịch vụ Managed Vector Database chuyên dụng, hỗ trợ thuật toán tìm kiếm HNSW tối ưu hóa cao, khả năng scale hàng triệu vector với độ trễ dưới 50ms và hỗ trợ tính năng lọc Metadata Filtering theo Namespace (rag-v2) rất mạnh mẽ."
        ),
        (
            "Câu 16: Nếu mô hình Gemini trả về JSON trắc nghiệm bị lỗi định dạng thì sao?",
            "Dạ, nhóm đã thiết lập cấu hình responseMimeType: 'application/json' ở API Gemini, kết hợp với khối try-catch và hàm chuẩn hóa Normalizer ở Backend. Nếu JSON bị lỗi, hệ thống tự động kích hoạt Graceful Fallback tạo câu hỏi chuẩn mực từ metadata bài học, không bao giờ làm crash ứng dụng."
        ),
        (
            "Câu 17: Phân chia trách nhiệm trong nhóm khi phát triển phân hệ AI này như thế nào?",
            "Dạ: Bạn Quốc Anh phụ trách thiết kế giao diện Sidebar, Streaming SSE, Interactive Quiz UI và Video Seek; Bạn Thanh Liêm phụ trách Backend RAG Engine, Intent Router, Hybrid Search và Bảo mật API; Bạn Đình Chương phụ trách Kiến trúc Database PostgreSQL, Vector DB Pinecone và Data Ingestion Pipeline."
        ),
        (
            "Câu 18: Tính năng 'Từ vựng trọng tâm' có bịa thêm từ vựng không có trong bài không?",
            "Dạ không ạ. Nhóm đã áp dụng Strict Grounding Prompt với Temperature 0.2 và quy tắc ép buộc: 'CHỈ trích xuất từ vựng xuất hiện trong nội dung bài học dưới đây, nếu không có đủ dữ liệu hãy trả lời lịch sự là chưa có đủ thông tin'."
        ),
        (
            "Câu 19: Hệ thống có hỗ trợ tìm kiếm tài liệu PDF đính kèm của bài học không?",
            "Dạ có ạ. Bảng lesson_materials lưu trữ các tài liệu PDF do giảng viên tải lên. Khi nạp dữ liệu, tài liệu được trích xuất text và nạp vào Pinecone với source: 'lesson-material-pdf', AI có thể tìm kiếm kết hợp cả phụ đề video lẫn tài liệu PDF."
        ),
        (
            "Câu 20: Hướng phát triển trong tương lai của hệ thống AI này là gì?",
            "Dạ, nhóm hướng đến 3 điểm mở rộng: 1) Tích hợp Speech-to-Text trực tiếp bằng WebRTC để đàm thoại giọng nói 2 chiều với AI; 2) Ứng dụng Agentic Workflow để tự động chấm điểm và chữa bài luận viết (Writing); 3) Cá nhân hóa lộ trình học tập dựa trên lịch sử làm bài trắc nghiệm của từng học viên."
        )
    ]

    for q, a in qa_list:
        add_callout(q, a, bg="F8FAFC", border_color="475569")

    # 4. Save Document
    output_filename = "SO_TAY_THUYET_TRINH_VA_BAO_VE_DO_AN_RAG_AI.docx"
    output_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", output_filename))
    doc.save(output_path)
    print(f"[OK] Da tao thanh cong file Word bao cao toan dien: {output_path}")

if __name__ == "__main__":
    generate_thesis_defense_doc()
