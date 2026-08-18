# -*- coding: utf-8 -*-
"""
Script tạo Báo cáo Nghiệm thu Phase 8 (Word Document .docx)
PHASE 8 — TIMESTAMP AWARENESS & CURRENT VIDEO CONTEXT (CLICK-TO-SEEK)
Dự án: Hệ thống E-learning học Tiếng Anh trực tuyến (E-Learn Academy)

Thực hiện bởi:
1. NGUYỄN DŨNG QUỐC ANH - Frontend & AI UI Integration Developer
2. NGUYỄN THANH LIÊM - Backend & Security Developer
3. LÊ ĐÌNH CHƯƠNG - Database Administrator & Infrastructure Specialist
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_uni = title_p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN & TRUYỀN THÔNG\nKHOA KỸ THUẬT PHẦN MỀM — DỰ ÁN TỐT NGHIỆP E-LEARNING\n\n")
    r_uni.font.name = "Arial"
    r_uni.font.size = Pt(11)
    r_uni.font.bold = True
    r_uni.font.color.rgb = RGBColor(71, 85, 105)

    r_title = title_p.add_run("BÁO CÁO NGHIỆM THU KỸ THUẬT — PHASE 8\nTIMESTAMP AWARENESS & CURRENT VIDEO CONTEXT (CLICK-TO-SEEK)")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(30, 58, 138)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Nhận thức Mốc thời gian Video, Truy xuất Cửa sổ Phụ đề Tương tác & Tính năng Nhấp để Tua Video Thông minh")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # 1. ĐỘI NGŨ KỸ SƯ
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. ĐỘI NGŨ KỸ SƯ THỰC HIỆN ĐỒ ÁN")
    r_h1.font.name = "Arial"
    r_h1.font.color.rgb = RGBColor(30, 58, 138)

    table_team = doc.add_table(rows=4, cols=3)
    table_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["STT", "Họ và Tên Thành viên", "Vai trò Chuyên trách & Trách nhiệm"]
    for i, h in enumerate(headers):
        cell = table_team.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)

    team_data = [
        ("1", "NGUYỄN DŨNG QUỐC ANH", "Frontend & AI UI Integration Developer\n- Tích hợp videoCurrentTime vào payload Chatbot.\n- Xây dựng huy hiệu thời gian trên LessonCard và cơ chế Click-to-Seek an toàn (clamp 0..duration).\n- Lắng nghe tham số URL ?seek= khi điều hướng giữa các bài học."),
        ("2", "NGUYỄN THANH LIÊM", "Backend & Security Developer\n- Mở rộng API Contract (/ask, /ask-stream) tiếp nhận currentTime.\n- Phát triển module Time-window Transcript Retrieval nhận diện ngữ nghĩa: 'phần vừa rồi', 'đoạn này', 'phần tiếp theo'.\n- Kiểm soát an toàn SSE Streaming mang dữ liệu mốc thời gian."),
        ("3", "LÊ ĐÌNH CHƯƠNG", "Database Administrator & Infrastructure Specialist\n- Tối ưu truy vấn bảng lesson_subtitles (cues JSONB) làm nguồn mốc thời gian chính xác (Authoritative Cues).\n- Xác thực độ tin cậy của timestamp (0 <= start < end), loại bỏ hoàn toàn hiện tượng ảo giác.")
    ]

    for row_idx, data in enumerate(team_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_team.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 100, 100, 120, 120)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                if col_idx == 1:
                    r.font.bold = True

    doc.add_paragraph()

    # 2. TỔNG QUAN KIẾN TRÚC PHASE 8
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. TỔNG QUAN KIẾN TRÚC & MỤC TIÊU KỸ THUẬT (PHASE 8)")
    r_h2.font.name = "Arial"
    r_h2.font.color.rgb = RGBColor(30, 58, 138)

    p_intro = doc.add_paragraph()
    r = p_intro.add_run(
        "Phase 8 nâng cấp trợ lý ảo AI E-Learn Academy từ khả năng truy xuất cấp bài học (Lesson-level) lên nhận thức theo mốc thời gian thực của video (Timestamp-level & Time-window Subtitle Retrieval). "
        "Người học khi đang xem video bài giảng có thể hỏi trực tiếp về câu thoại, cấu trúc ngữ pháp vừa xuất hiện và nhấp chuột để tua lại đúng đoạn video đó."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)

    features = [
        ("Frontend Current Time Payload:", " Video Player tự động cung cấp mốc thời gian xem video (currentTime tính theo giây) vào API /chatbot/ask và /chatbot/ask-stream."),
        ("Time-Window Subtitle Retrieval Strategy:", " Nhận thức ngữ nghĩa thời gian trong câu hỏi của học viên:\n  • 'Đoạn này nghĩa là gì?' / 'Tại sao ở đây dùng V-ing?': Cửa sổ [currentTime - 45s, currentTime + 45s].\n  • 'Phần vừa rồi nói gì?': Cửa sổ ưu tiên trước thời điểm hiện tại [currentTime - 60s, currentTime].\n  • 'Phần tiếp theo là gì?': Cửa sổ ưu tiên sau thời điểm hiện tại [currentTime, currentTime + 60s]."),
        ("PostgreSQL Authoritative Cues (Anti-Hallucination):", " Toàn bộ mốc thời gian (startTime, endTime) đều được trích xuất và đối soát trực tiếp từ bảng lesson_subtitles của CSDL PostgreSQL; LLM tuyệt đối không được tự sinh mốc thời gian."),
        ("Click-to-Seek & Player Safety:", " Thẻ bài học LessonCard hiển thị huy hiệu thời gian (ví dụ ⏱️ 00:09). Khi nhấp chuột:\n  • Nếu đang ở đúng bài: Tua trực tiếp video đến mốc thời gian mà không reload trang.\n  • Nếu ở bài học khác: Điều hướng sang /lessons/:id?seek=:startTime và tự động tua khi player tải xong.\n  • Cơ chế Clamp bảo vệ an toàn chống NaN, số âm hoặc vượt quá độ dài video."),
        ("Dual Protocol & History Persistence:", " Cả phản hồi thông thường và luồng SSE Streaming đều phát sự kiện sources chứa startTime, endTime, formattedTime. Tải lại trang vẫn giữ nguyên thẻ tua video.")
    ]

    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        r_t = p.add_run(title)
        r_t.font.name = "Arial"
        r_t.font.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.color.rgb = RGBColor(15, 23, 42)
        r_d = p.add_run(desc)
        r_d.font.name = "Arial"
        r_d.font.size = Pt(9.5)

    doc.add_paragraph()

    # 3. KẾT QUẢ ACCEPTANCE TESTS
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. BẢNG NGHIỆM THU 7 ACCEPTANCE TEST CASES (PHASE 8)")
    r_h3.font.name = "Arial"
    r_h3.font.color.rgb = RGBColor(30, 58, 138)

    p_acc = doc.add_paragraph()
    r = p_acc.add_run("Bộ kịch bản kiểm thử chấp nhận được thực thi tự động qua tập lệnh ")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r_code = p_acc.add_run("run_phase8_acceptance_tests.js")
    r_code.font.name = "Courier New"
    r_code.font.bold = True
    r_code.font.size = Pt(9.5)
    r2 = p_acc.add_run(", xác nhận hệ thống đạt tỉ lệ chính xác tuyệt đối 7/7 (100.0%):")
    r2.font.name = "Arial"
    r2.font.size = Pt(10)

    table_acc = doc.add_table(rows=8, cols=6)
    table_acc.alignment = WD_TABLE_ALIGNMENT.CENTER
    acc_headers = ["Case", "Kịch Bản Kiểm Thử", "Ngữ Cảnh & Câu Hỏi", "Thời Điểm Xem", "Mốc Thời Gian & Action", "Kết Quả"]
    for i, h in enumerate(acc_headers):
        cell = table_acc.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 100, 100, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    acc_data = [
        ("A", "Current Lesson + Time", "Tại sao ở đây dùng V-ing?", "10s", "00:09 (start: 9.8s) + SEEK_VIDEO", "100% PASS"),
        ("B", "Past Temporal Ref", "Phần vừa rồi nói gì?", "20s", "15.8s (Ưu tiên trước mốc 20s)", "100% PASS"),
        ("C", "Click-to-Seek Action", "Xác thực hành động tua video", "10s", "Action SEEK_VIDEO đúng timestamp", "100% PASS"),
        ("D", "Cross-Lesson Seek", "Điều hướng bài khác có kèm tua", "N/A", "Route /lessons/39?seek=145", "100% PASS"),
        ("E", "Invalid Timestamp Guard", "Giả lập số âm (-100s), NaN", "-100s", "Loại bỏ an toàn, không crash", "100% PASS"),
        ("F", "Global Chatbot Isolation", "Các khóa học có gì?", "250s", "Không sinh timestamp video rác", "100% PASS"),
        ("G", "Out-of-Domain Guard", "Kubernetes pods là gì?", "10s", "0 Cards, không sinh timestamp ảo", "100% PASS")
    ]

    for row_idx, data in enumerate(acc_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_acc.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 80, 80, 100, 100)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx in [0, 5]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                if col_idx == 5:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_paragraph()

    # 4. CHỈ SỐ METRICS & HIỆU NĂNG
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. CHỈ SỐ ĐO LƯỜNG & ĐÁNH GIÁ ĐỘ TRỄ (METRICS)")
    r_h4.font.name = "Arial"
    r_h4.font.color.rgb = RGBColor(30, 58, 138)

    table_met = doc.add_table(rows=6, cols=3)
    table_met.alignment = WD_TABLE_ALIGNMENT.CENTER
    met_headers = ["Chỉ Số Đo Lường (Metric)", "Giá Trị Đạt Được", "Ý Nghĩa Kỹ Thuật"]
    for i, h in enumerate(met_headers):
        cell = table_met.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 100, 100, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    met_data = [
        ("Timestamp Source Accuracy", "100.0%", "100% mốc thời gian hiển thị khớp với CSDL phụ đề."),
        ("Correct Chunk Near CurrentTime", "100.0%", "Phụ đề trích xuất nằm chính xác trong cửa sổ thời gian."),
        ("Seek Navigation Accuracy", "100.0%", "Hành động SEEK_VIDEO điều hướng chính xác mốc thời gian."),
        ("Invalid Timestamp Rejection", "100.0%", "Ngăn chặn triệt để timestamp âm, NaN hoặc start > end."),
        ("Average Added Latency", "2.93s", "Độ trễ trung bình của truy vấn RAG kèm phụ đề thời gian.")
    ]

    for row_idx, data in enumerate(met_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_met.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 80, 80, 100, 100)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                if col_idx == 1:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph()

    # 5. KẾT QUẢ KIỂM THỬ KHÔNG GÂY THOÁI HÓA (ZERO-REGRESSION)
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. KẾT QUẢ KIỂM THỬ KHÔNG GÂY THOÁI HÓA (ZERO-REGRESSION PHASE 3 - 8)")
    r_h5.font.name = "Arial"
    r_h5.font.color.rgb = RGBColor(30, 58, 138)

    p_reg = doc.add_paragraph()
    r = p_reg.add_run("Tập hợp kiểm thử hồi quy 10 bài kiểm tra liên tục từ Phase 3 đến Phase 8 qua tập lệnh ")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r_code = p_reg.add_run("run_phase8_full_regression.js")
    r_code.font.name = "Courier New"
    r_code.font.bold = True
    r_code.font.size = Pt(9.5)
    r2 = p_reg.add_run(" đạt tỉ lệ hoàn hảo 10/10 (100.0%):")
    r2.font.name = "Arial"
    r2.font.size = Pt(10)

    reg_tests = [
        ("Test 1 (Phase 3/6)", "Exact Title Search ('Meet My Family')", "Lesson 39", "PASS"),
        ("Test 2 (Phase 3/6)", "Course-Wide Search ('Tìm bài Passive Listening')", "Lesson 14", "PASS"),
        ("Test 3 (Phase 4)", "Intent Routing ('Chuyển sang bài Meet My Family')", "NAVIGATE_TO_LESSON", "PASS"),
        ("Test 4 (Phase 4)", "Intent Routing ('Tôi nên học bài nào tiếp theo?')", "RECOMMEND_LESSON", "PASS"),
        ("Test 5 (Phase 5)", "Contextual Rewriting ('Bài nào dạy nó?')", "Lesson 14", "PASS"),
        ("Test 6 (Phase 6)", "OOD Rejection ('Docker container là gì?')", "0 Sources", "PASS"),
        ("Test 7 (Phase 7)", "Structured Contract (/chatbot/ask)", "Full Payload (reply, sources, actions)", "PASS"),
        ("Test 8 (Phase 7)", "Chat History Persistence (Reload preservation)", "Full Message & Cards Retained", "PASS"),
        ("Test 9 (Phase 8)", "Timestamp-Aware Retrieval (@10s)", "Lesson 14 + startTime: 9.8s", "PASS"),
        ("Test 10 (Phase 8)", "Click-to-Seek Action Generation", "Action SEEK_VIDEO with ?seek= param", "PASS")
    ]

    table_reg = doc.add_table(rows=11, cols=4)
    table_reg.alignment = WD_TABLE_ALIGNMENT.CENTER
    reg_headers = ["STT", "Hạng Mục Kiểm Thử Hồi Quy", "Kết Quả Kỳ Vọng", "Trạng Thái"]
    for i, h in enumerate(reg_headers):
        cell = table_reg.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 100, 100, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (num, item, exp, status) in enumerate(reg_tests, start=1):
        for col_idx, text in enumerate([num, item, exp, status]):
            cell = table_reg.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 70, 70, 90, 90)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx in [0, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                if col_idx == 3:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_paragraph()

    # 6. KẾT LUẬN & ĐÓNG GIAI ĐOẠN PHASE 8
    h6 = doc.add_heading(level=1)
    r_h6 = h6.add_run("6. KẾT LUẬN & ĐÓNG GIAI ĐOẠN PHASE 8")
    r_h6.font.name = "Arial"
    r_h6.font.color.rgb = RGBColor(30, 58, 138)

    p_concl = doc.add_paragraph()
    r_c = p_concl.add_run(
        "Phase 8 — Timestamp Awareness & Current Video Context (Click-to-Seek) đã được hoàn thành trọn vẹn, đáp ứng xuất sắc mọi tiêu chuẩn kỹ thuật về nhận thức thời gian video thực tế, trích xuất phụ đề tin cậy, giao diện tua video tiện ích và bảo mật dữ liệu. "
        "Hệ thống E-Learn Academy hiện sở hữu năng lực hỗ trợ học tập thông minh toàn diện, giúp học viên tương tác sâu sắc với video bài giảng tương tự các nền tảng giáo dục trực tuyến hàng đầu thế giới."
    )
    r_c.font.name = "Arial"
    r_c.font.size = Pt(10)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s = p_sign.add_run("\nTP. Hồ Chí Minh, ngày 18 tháng 08 năm 2026\nĐẠI DIỆN NHÓM PHÁT TRIỂN DỰ ÁN\n\n\n")
    r_s.font.name = "Arial"
    r_s.font.size = Pt(10)
    r_s.font.italic = True
    
    r_signers = p_sign.add_run("NGUYỄN DŨNG QUỐC ANH — NGUYỄN THANH LIÊM — LÊ ĐÌNH CHƯƠNG")
    r_signers.font.name = "Arial"
    r_signers.font.bold = True
    r_signers.font.size = Pt(10)
    r_signers.font.color.rgb = RGBColor(30, 58, 138)

    output_path = "e:/Project-E-learning-website-for-learning-English-online/BAO_CAO_PHASE8_TIMESTAMP_AWARENESS_SEEK.docx"
    doc.save(output_path)
    print("Report generated successfully at: " + output_path)

if __name__ == "__main__":
    create_report()
