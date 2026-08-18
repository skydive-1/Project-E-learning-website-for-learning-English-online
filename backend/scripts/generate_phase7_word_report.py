# -*- coding: utf-8 -*-
"""
Script tạo Báo cáo Nghiệm thu Phase 7 (Word Document .docx)
PHASE 7 — STRUCTURED SOURCES, LESSON CARDS & NAVIGATION
Dự án: Hệ thống E-learning học Tiếng Anh trực tuyến (E-Learn Academy)

Thực hiện bởi:
1. NGUYỄN DŨNG QUỐC ANH - Frontend & AI UI Integration Developer
2. NGUYỄN THANH LIÊM - Backend & Security Developer
3. LÊ ĐÌNH CHƯƠNG - Database Administrator & Infrastructure Specialist
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Đặt màu nền cho cell trong bảng"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt padding cho cell"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()

    # Thiết lập lề trang chuẩn A4
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # Style Tiêu đề chính
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_uni = title_p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN & TRUYỀN THÔNG\nKHOA KỸ THUẬT PHẦN MỀM — DỰ ÁN TỐT NGHIỆP E-LEARNING\n\n")
    r_uni.font.name = "Arial"
    r_uni.font.size = Pt(11)
    r_uni.font.bold = True
    r_uni.font.color.rgb = RGBColor(71, 85, 105)

    r_title = title_p.add_run("BÁO CÁO NGHIỆM THU KỸ THUẬT — PHASE 7\nSTRUCTURED SOURCES, LESSON CARDS & NAVIGATION")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(30, 58, 138)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Kiến trúc Phản hồi Có cấu trúc, Xác thực Nguồn Dữ liệu PostgreSQL & Giao diện Thẻ Bài học Chuẩn Udemy")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # ==================== THÔNG TIN NHÓM ====================
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. ĐỘI NGŨ KỸ SƯ THỰC HIỆN ĐỒ ÁN")
    r_h1.font.name = "Arial"
    r_h1.font.color.rgb = RGBColor(30, 58, 138)

    table_team = doc.add_table(rows=4, cols=3)
    table_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["STT", "Họ và Tên Thành viên", "Vai trò Chuyên trách & Trách nhiệm"]
    for i, h in enumerate(headers):
        cell = table_team.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)

    team_data = [
        ("1", "NGUYỄN DŨNG QUỐC ANH", "Frontend & AI UI Integration Developer\n- Thiết kế & xây dựng UI LessonCard (Udemy UX style).\n- Xử lý SSE Event Streaming (metadata, tokens, sources) và điều hướng React Router."),
        ("2", "NGUYỄN THANH LIÊM", "Backend & Security Developer\n- Xây dựng API Contract (/ask, /ask-stream) phản hồi có cấu trúc.\n- Thiết lập kiểm soát quyền truy cập khóa học và bộ lọc Guardrails."),
        ("3", "LÊ ĐÌNH CHƯƠNG", "Database Administrator & Infrastructure Specialist\n- Xây dựng PostgreSQL Authoritative Source of Truth Service.\n- Thiết kế cơ chế lưu trữ lịch sử hội thoại có cấu trúc không làm phình schema.")
    ]

    for row_idx, data in enumerate(team_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_team.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 100, 100, 120, 120)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                if col_idx == 1:
                    r.font.bold = True

    doc.add_paragraph()

    # ==================== TỔNG QUAN PHASE 7 ====================
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. TỔNG QUAN KIẾN TRÚC & MỤC TIÊU UX (PHASE 7)")
    r_h2.font.name = "Arial"
    r_h2.font.color.rgb = RGBColor(30, 58, 138)

    p_intro = doc.add_paragraph()
    r = p_intro.add_run(
        "Tiếp nối thành công của các Phase 3-6 (Metadata V2, Dynamic Scope Routing, Conversational Query Rewriting, Hybrid Search & Deterministic Reranking), "
        "Phase 7 tập trung hoàn thiện tầng trình bày và trải nghiệm người dùng tương đương hệ thống trợ lý học tập chuẩn quốc tế (Udemy AI Assistant)."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)

    features = [
        ("Structured Backend Response:", " Phản hồi của AI được đóng gói thành schema JSON chặt chẽ gồm: intent, reply, sources, actions. Tách biệt hoàn toàn giữa nội dung văn bản và siêu dữ liệu điều hướng, loại bỏ triệt để việc parse regex text thiếu tin cậy."),
        ("PostgreSQL Authoritative Source of Truth (Anti-Hallucination):", " LLM tuyệt đối không được tự bịa Lesson ID hay tiêu đề bài học. 100% thẻ bài học trả về cho người học đều được đối chiếu và xác thực tồn tại trực tiếp từ CSDL PostgreSQL của khóa học."),
        ("Intent-Specific Source Behavior:", " Trả về nguồn tương ứng với từng Intent nghiệp vụ:\n  • SEARCH_LESSON / COURSE_QA: Trả về Top 1–3 bài học phù hợp nhất từ Hybrid Reranker.\n  • NAVIGATE_TO_LESSON: Trả về đúng 1 thẻ bài học đích cần mở kèm action OPEN_LESSON.\n  • RECOMMEND_LESSON: Trả về 1–3 thẻ đề xuất học tiếp theo lộ trình khóa học (badge 'Đề xuất học tiếp').\n  • GENERAL_ENGLISH_QA / Out-of-Domain: sources = [] (tuyệt đối không hiển thị thẻ bài học rác)."),
        ("Dual Protocol Alignment (Streaming & Non-Streaming):", " Đồng nhất cấu trúc dữ liệu trên cả 2 kênh giao tiếp: REST API thông thường (/chatbot/ask) và giao thức truyền tải dòng Server-Sent Events (/chatbot/ask-stream) với các sự kiện metadata, token, sources, done."),
        ("Frontend Lesson Card & Seamless Navigation:", " Component LessonCard được thiết kế hiện đại (Glassmorphism, micro-animations, badges) và điều hướng tự động vào đúng Route thực tế của hệ thống (/lessons/:lessonId) khi người học nhấp chuột."),
        ("Zero-Migration Conversation Persistence:", " Lưu trữ lịch sử tin nhắn kèm thẻ bài học xác thực vào bảng ai_chat hiện có, đảm bảo khi tải lại trang (reload) toàn bộ Lesson Cards vẫn được tái hiện nguyên vẹn.")
    ]

    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        r_t = p.add_run(title)
        r_t.font.name = "Arial"
        r_t.font.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.color.rgb = RGBColor(15, 23, 42)
        r_d = p.add_run(desc)
        r_d.font.name = "Arial"
        r_d.font.size = Pt(9.5)

    doc.add_paragraph()

    # ==================== KẾT QUẢ ACCEPTANCE TESTS ====================
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. BẢNG NGHIỆM THU TOÀN DIỆN 7 ACCEPTANCE TEST CASES")
    r_h3.font.name = "Arial"
    r_h3.font.color.rgb = RGBColor(30, 58, 138)

    p_acc = doc.add_paragraph()
    r = p_acc.add_run("Bộ kịch bản kiểm thử chấp nhận (Acceptance Tests) được thực thi độc lập và tự động hóa qua tập lệnh ")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r_code = p_acc.add_run("run_phase7_acceptance_tests.js")
    r_code.font.name = "Courier New"
    r_code.font.bold = True
    r_code.font.size = Pt(9.5)
    r2 = p_acc.add_run(", xác nhận hệ thống đạt tỉ lệ chính xác tuyệt đối 7/7 (100.0%):")
    r2.font.name = "Arial"
    r2.font.size = Pt(10)

    table_acc = doc.add_table(rows=8, cols=6)
    table_acc.alignment = WD_TABLE_ALIGNMENT.CENTER
    acc_headers = ["Case", "Loại Nghiệp Vụ", "Câu Hỏi Đầu Vào", "Intent Xác Định", "Số Lượng Nguồn & Bài Đích", "Kết Quả"]
    for i, h in enumerate(acc_headers):
        cell = table_acc.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 100, 100, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    acc_data = [
        ("1", "Search Lesson", "Bài nào dạy Passive Listening?", "SEARCH_LESSON", "1 Card (Lesson 14: Passive Listening)", "100% PASS"),
        ("2", "Semantic Search", "Bài nào dạy phương pháp nghe khi đang làm việc khác?", "SEARCH_LESSON", "1 Card (Lesson 14: Passive Listening)", "100% PASS"),
        ("3", "Navigate Intent", "Đưa tôi tới bài Meet My Family.", "NAVIGATE_TO_LESSON", "1 Target Card (Lesson 39) + OPEN_LESSON", "100% PASS"),
        ("4", "Recommend Intent", "Học xong bài này tôi nên học bài nào?", "RECOMMEND_LESSON", "2 Cards (Đề xuất học tiếp lộ trình)", "100% PASS"),
        ("5", "Contextual Follow-up", "Bài nào nói về nó? (sau Passive Listening)", "SEARCH_LESSON", "1 Card (Lesson 14 resolved qua Rewriter)", "100% PASS"),
        ("6", "Out-of-Domain Rejection", "Bài nào dạy Kubernetes?", "SEARCH_LESSON", "0 Cards (Rejection dưới threshold)", "100% PASS"),
        ("7", "General English QA", "Hello nghĩa là gì?", "GENERAL_ENGLISH_QA", "0 Cards (Không tạo card bài học rác)", "100% PASS")
    ]

    for row_idx, data in enumerate(acc_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_acc.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 80, 80, 100, 100)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx in [0, 5]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                if col_idx == 5:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_paragraph()

    # ==================== BẢO MẬT & INTEGRITY ====================
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. KIỂM TRA BẢO MẬT & TOÀN VẸN DỮ LIỆU CƠ SỞ DỮ LIỆU")
    r_h4.font.name = "Arial"
    r_h4.font.color.rgb = RGBColor(30, 58, 138)

    sec_items = [
        ("Chống sinh dữ liệu ảo (Anti-Hallucination Guard):", " Khi cố ý giả lập đầu vào với lessonId không tồn tại (ID 999999), module Source Builder lập tức loại bỏ khỏi danh sách kết quả, đảm bảo người dùng không bao giờ nhìn thấy thẻ bài học hỏng."),
        ("Cách ly ranh giới khóa học (Course Boundary Isolation):", " Khi truy vấn bài học thuộc Course khác (ví dụ Lesson 39 của Course 22) trong ngữ cảnh truy xuất của Course 5, hệ thống từ chối triệt để 100%, bảo vệ toàn vẹn phân quyền dữ liệu giữa các khóa học."),
        ("Định dạng SSE Streaming chuẩn hóa:", " Giao thức /chatbot/ask-stream phát đúng chuỗi sự kiện: Metadata Event (khởi tạo intent/scope) -> Token Events (truyền dòng phản hồi) -> Sources Event (danh sách thẻ bài học) -> [DONE] Event.")
    ]

    for title, desc in sec_items:
        p = doc.add_paragraph(style='List Bullet')
        r_t = p.add_run(title)
        r_t.font.name = "Arial"
        r_t.font.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.color.rgb = RGBColor(15, 23, 42)
        r_d = p.add_run(desc)
        r_d.font.name = "Arial"
        r_d.font.size = Pt(9.5)

    doc.add_paragraph()

    # ==================== KẾT QUẢ ZERO-REGRESSION ====================
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. KẾT QUẢ KIỂM THỬ KHÔNG GÂY THOÁI HÓA (ZERO-REGRESSION)")
    r_h5.font.name = "Arial"
    r_h5.font.color.rgb = RGBColor(30, 58, 138)

    p_reg = doc.add_paragraph()
    r = p_reg.add_run("Tập hợp kiểm thử hồi quy 8 bài kiểm tra từ Phase 3 đến Phase 7 được thực thi qua tập lệnh ")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r_code = p_reg.add_run("run_phase7_full_regression.js")
    r_code.font.name = "Courier New"
    r_code.font.bold = True
    r_code.font.size = Pt(9.5)
    r2 = p_reg.add_run(" đạt tỉ lệ hoàn hảo 8/8 (100.0%):")
    r2.font.name = "Arial"
    r2.font.size = Pt(10)

    reg_tests = [
        ("Test 1 (Phase 3/6)", "Exact Title Search ('Meet My Family')", "Lesson 39", "PASS"),
        ("Test 2 (Phase 3/6)", "Course-Wide Search ('Tìm bài Passive Listening')", "Lesson 14", "PASS"),
        ("Test 3 (Phase 4)", "Intent Routing ('Chuyển sang bài Meet My Family')", "NAVIGATE_TO_LESSON", "PASS"),
        ("Test 4 (Phase 4)", "Intent Routing ('Tôi nên học bài nào tiếp theo?')", "RECOMMEND_LESSON", "PASS"),
        ("Test 5 (Phase 5)", "Contextual Rewriting ('Bài nào dạy nó?')", "Lesson 14", "PASS"),
        ("Test 6 (Phase 6)", "OOD Rejection ('Docker container là gì?')", "0 Sources", "PASS"),
        ("Test 7 (Phase 7)", "Structured Contract (/chatbot/ask)", "Full Payload (reply, sources, actions)", "PASS"),
        ("Test 8 (Phase 7)", "Chat History Persistence (Reload preservation)", "Full Message & Cards Retained", "PASS")
    ]

    table_reg = doc.add_table(rows=9, cols=4)
    table_reg.alignment = WD_TABLE_ALIGNMENT.CENTER
    reg_headers = ["STT", "Hạng Mục Kiểm Thử Hồi Quy", "Kết Quả Kỳ Vọng", "Trạng Thái"]
    for i, h in enumerate(reg_headers):
        cell = table_reg.cell(0, i)
        cell.text = h
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 100, 100, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Arial"
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (num, item, exp, status) in enumerate(reg_tests, start=1):
        for col_idx, text in enumerate([num, item, exp, status]):
            cell = table_reg.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, 70, 70, 90, 90)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            if col_idx in [0, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                if col_idx == 3:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_paragraph()

    # ==================== KẾT LUẬN & ĐÓNG PHASE 7 ====================
    h6 = doc.add_heading(level=1)
    r_h6 = h6.add_run("6. KẾT LUẬN & ĐÓNG GIAI ĐOẠN PHASE 7")
    r_h6.font.name = "Arial"
    r_h6.font.color.rgb = RGBColor(30, 58, 138)

    p_concl = doc.add_paragraph()
    r_c = p_concl.add_run(
        "Phase 7 — Structured Sources, Lesson Cards & Navigation đã được triển khai hoàn tất 100% các tiêu chí kỹ thuật và trải nghiệm người dùng theo đúng hợp đồng thiết kế. "
        "Hệ thống E-learning AI Assistant hiện tại sở hữu đầy đủ chuỗi giá trị từ Truy xuất Đa phương thức (Hybrid Retrieval), Định tuyến Ý định (Intent Router), Viết lại Truy vấn (Query Rewriter), Xác thực Nguồn Dữ liệu (PostgreSQL Authoritative Sources), đến Giao diện Thẻ Bài học & Điều hướng Trực quan (Lesson Cards UI)."
    )
    r_c.font.name = "Arial"
    r_c.font.size = Pt(10)

    # Chữ ký xác nhận
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s = p_sign.add_run("\nTP. Hồ Chí Minh, ngày 18 tháng 08 năm 2026\nĐẠI DIỆN NHÓM PHÁT TRIỂN DỰ ÁN\n\n\n")
    r_s.font.name = "Arial"
    r_s.font.size = Pt(10)
    r_s.font.italic = True
    
    r_signers = p_sign.add_run("NGUYỄN DŨNG QUỐC ANH — NGUYỄN THANH LIÊM — LÊ ĐÌNH CHƯƠNG")
    r_signers.font.name = "Arial"
    r_signers.font.bold = True
    r_signers.font.size = Pt(10)
    r_signers.font.color.rgb = RGBColor(30, 58, 138)

    output_path = "e:/Project-E-learning-website-for-learning-English-online/BAO_CAO_PHASE7_STRUCTURED_SOURCES_LESSON_CARDS.docx"
    doc.save(output_path)
    print("Report generated successfully at: " + output_path)

if __name__ == "__main__":
    create_report()
