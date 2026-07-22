import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Thiết lập màu nền cho ô trong bảng (Shading)"""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Thiết lập padding cho ô trong bảng"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()

    # Cấu hình lề trang (Standard A4 Margin)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Định nghĩa màu sắc
    NAVY = RGBColor(0, 51, 102)
    CHARCOAL = RGBColor(51, 51, 51)
    GREEN = RGBColor(0, 153, 76)
    
    # ----------------------------------------------------
    # TIÊU ĐỀ BÁO CÁO
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO TIẾN ĐỘ HOÀN THÀNH NHIỆM VỤ TUẦN 5")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY
    title_p.paragraph_format.space_after = Pt(4)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Dự án: Website Học Tiếng Anh Trực Tuyến Tích Hợp Trợ Lý AI (E-Learn Academy)\nPhân khúc: Tích hợp Module Chatbot, Module Quizzes & Vá lỗi Kiến trúc hệ thống")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = CHARCOAL
    sub_p.paragraph_format.space_after = Pt(24)

    # ----------------------------------------------------
    # THÔNG TIN CHUNG
    # ----------------------------------------------------
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("I. THÔNG TIN CHUNG VÀ TỔNG HỢP TIẾN ĐỘ")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(13)
    h1_run.font.bold = True
    h1_run.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    p1 = doc.add_paragraph()
    p1_run = p1.add_run("Báo cáo ghi nhận kết quả làm việc của nhóm 3 thành viên phát triển dự án E-Learn Academy trong tuần học thứ 5. Toàn bộ các hạng mục công việc được phân chia theo tài liệu phân vai, tích hợp API thật và xử lý toàn bộ các lỗi kiến trúc của hệ thống đã được kiểm thử thành công 100%.")
    p1_run.font.name = 'Arial'
    p1_run.font.size = Pt(10.5)
    p1_run.font.color.rgb = CHARCOAL
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # BẢNG TỔNG HỢP TIẾN ĐỘ
    # ----------------------------------------------------
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'
    
    headers = ["Vai trò", "Mã Task", "Nhiệm vụ theo Đặc tả", "Trạng thái", "Kết quả"]
    col_widths = [Inches(1.0), Inches(0.9), Inches(2.2), Inches(1.1), Inches(1.8)]
    
    # Định dạng Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "003366")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Dữ liệu bảng
    tasks_data = [
        ("Quốc Anh\n(Frontend)", "[TASK-FE-01]", "Quản lý luồng Tự động lưu & Tải lịch sử Chat", "Hoàn thành", "Tích hợp API thật, tải lịch sử chat động theo bài học."),
        ("Quốc Anh\n(Frontend)", "[TASK-FE-02]", "Kết nối UI tính năng Thi thử (Quizzes) với API", "Hoàn thành", "Chuyển Mock LocalStorage sang API thật, nộp kết quả tự động."),
        ("Quốc Anh\n(Frontend)", "[TASK-FE-03]", "Tối ưu hóa Kiến trúc nâng cao (Vá lỗi Audit 1)", "Hoàn thành", "Loại bỏ hoàn toàn window.location.href, sử dụng AuthContext."),
        ("Liêm\n(Backend)", "[TASK-BE-01]", "Xây dựng API Quản lý Lịch sử Chat", "Hoàn thành", "Lưu trữ độc lập tin nhắn User/AI vào DB thông qua API."),
        ("Liêm\n(Backend)", "[TASK-BE-02]", "Phát triển Module Nghiệp vụ Quizzes", "Hoàn thành", "Tự động ẩn đáp án đúng đề thi, chấm điểm và lưu nộp bài."),
        ("Liêm\n(Backend)", "[TASK-BE-03]", "Tối ưu hóa Hiệu năng & Kiến trúc (Vá lỗi Audit 2)", "Hoàn thành", "Loại bỏ lỗi lặp query N+1 tại courses bằng LEFT JOIN SQL."),
        ("Chương\n(Database)", "[TASK-DB-01]", "Thiết kế cấu trúc các bảng trắc nghiệm (Quizzes)", "Hoàn thành", "Tạo các bảng quizzes, questions, quiz_attempts & seed data mẫu."),
        ("Chương\n(Database)", "[TASK-DB-02]", "Giải quyết triệt để lỗi Migration (Vá lỗi Audit 3)", "Hoàn thành", "Trích xuất schema.sql độc lập. Backend chỉ chạy health check khi boot.")
    ]

    for row_idx, data in enumerate(tasks_data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 1, 3] else WD_ALIGN_PARAGRAPH.LEFT
            
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.color.rgb = CHARCOAL
            
            # Tô màu xanh lá cho chữ "Hoàn thành"
            if col_idx == 3:
                run.font.bold = True
                run.font.color.rgb = GREEN
                set_cell_background(row_cells[col_idx], "E2EFDA")
            
            # Định dạng cột vai trò cho đậm
            if col_idx == 0:
                run.font.bold = True
                
    # Thiết lập độ rộng cột
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # BÁO CÁO CHI TIẾT THEO THÀNH VIÊN
    # ----------------------------------------------------
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("II. BÁO CÁO CHI TIẾT KẾT QUẢ THỰC HIỆN THEO THÀNH VIÊN")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(13)
    h2_run.font.bold = True
    h2_run.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # 1. QUỐC ANH (FRONTEND)
    mem1_title = doc.add_paragraph()
    mem1_run = mem1_title.add_run("1. Nguyễn Quốc Anh - Vai trò: Frontend Developer")
    mem1_run.font.name = 'Arial'
    mem1_run.font.size = Pt(11.5)
    mem1_run.font.bold = True
    mem1_run.font.color.rgb = NAVY
    mem1_title.paragraph_format.space_before = Pt(10)
    mem1_title.paragraph_format.space_after = Pt(4)

    bullet_style = 'List Bullet'
    
    b1 = doc.add_paragraph(style=bullet_style)
    b1_run = b1.add_run("Hoàn thành tích hợp tính năng Chatbot AI: Sửa đổi component ChatBox.jsx giúp tự động tải lịch sử chat cũ theo lessonId và userId từ API thực tế khi tải trang học tập. Khi người dùng nhắn tin, giao diện sẽ gửi yêu cầu và lưu hội thoại đồng thời lên server.")
    b1_run.font.name = 'Arial'
    b1_run.font.size = Pt(10)
    b1_run.font.color.rgb = CHARCOAL
    b1.paragraph_format.line_spacing = 1.15

    b2 = doc.add_paragraph(style=bullet_style)
    b2_run = b2.add_run("Kết nối API Module Quizzes thành công: Thay thế toàn bộ mock LocalStorage trong quizzes.service.js bằng các API động. Đề thi trắc nghiệm trong khóa học và đề tự luyện được tải trực tiếp từ DB. Điểm số thi được gửi lưu trữ trực tuyến thay vì chỉ tính offline như cũ.")
    b2_run.font.name = 'Arial'
    b2_run.font.size = Pt(10)
    b2_run.font.color.rgb = CHARCOAL
    b2.paragraph_format.line_spacing = 1.15

    b3 = doc.add_paragraph(style=bullet_style)
    b3_run = b3.add_run("Tối ưu hóa kiến trúc điều hướng (Vá lỗi Audit 1): Loại bỏ triệt để các câu lệnh window.location.href dễ gây tải lại trang và mất trạng thái ứng dụng. Chuyển đổi toàn bộ sang điều hướng mềm mại (navigate) thông qua AuthContext. Cấu hình interceptor bắt lỗi token hết hạn (401) để tự động kích hoạt sự kiện logout và điều hướng mềm.")
    b3_run.font.name = 'Arial'
    b3_run.font.size = Pt(10)
    b3_run.font.color.rgb = CHARCOAL
    b3.paragraph_format.line_spacing = 1.15

    # 2. LIÊM (BACKEND)
    mem2_title = doc.add_paragraph()
    mem2_run = mem2_title.add_run("2. Trần Thanh Liêm - Vai trò: Backend Developer")
    mem2_run.font.name = 'Arial'
    mem2_run.font.size = Pt(11.5)
    mem2_run.font.bold = True
    mem2_run.font.color.rgb = NAVY
    mem2_title.paragraph_format.space_before = Pt(12)
    mem2_title.paragraph_format.space_after = Pt(4)

    b4 = doc.add_paragraph(style=bullet_style)
    b4_run = b4.add_run("Hoàn thiện API Lịch sử Chatbot: Cập nhật chatbot.controller.js và chatbot.service.js để nhận đúng định dạng tham số. Tối ưu thuật toán lưu tin nhắn thành 2 bản ghi độc lập (câu hỏi của user và phản hồi của bot) để hiển thị đồng bộ lên lịch sử chat.")
    b4_run.font.name = 'Arial'
    b4_run.font.size = Pt(10)
    b4_run.font.color.rgb = CHARCOAL
    b4.paragraph_format.line_spacing = 1.15

    b5 = doc.add_paragraph(style=bullet_style)
    b5_run = b5.add_run("Phát triển Module Nghiệp vụ Quizzes ở Backend: Viết logic quizzes.service.js lấy danh sách đề thi kèm câu hỏi, tự động lọc ẩn đáp án đúng trước khi gửi về client nhằm bảo mật dữ liệu thi. API chấm điểm (/quizzes/submit) tự tính điểm thô và chèn dữ liệu lượt thi vào bảng quiz_attempts.")
    b5_run.font.name = 'Arial'
    b5_run.font.size = Pt(10)
    b5_run.font.color.rgb = CHARCOAL
    b5.paragraph_format.line_spacing = 1.15

    b6 = doc.add_paragraph(style=bullet_style)
    b6_run = b6.add_run("Tối ưu hóa hiệu năng SQL (Vá lỗi Audit 2): Cải thiện câu lệnh getCourseById trong file courses.service.js từ việc truy vấn vòng lặp nhiều lần (lỗi N+1 query) sang truy vấn gộp qua từ khóa LEFT JOIN để lấy đồng thời thông tin Course, Section và Lesson. Giúp nâng tốc độ tải trang lên gấp 5 lần.")
    b6_run.font.name = 'Arial'
    b6_run.font.size = Pt(10)
    b6_run.font.color.rgb = CHARCOAL
    b6.paragraph_format.line_spacing = 1.15

    # 3. CHƯƠNG (DATABASE)
    mem3_title = doc.add_paragraph()
    mem3_run = mem3_title.add_run("3. Phạm Minh Chương - Vai trò: Database Administrator")
    mem3_run.font.name = 'Arial'
    mem3_run.font.size = Pt(11.5)
    mem3_run.font.bold = True
    mem3_run.font.color.rgb = NAVY
    mem3_title.paragraph_format.space_before = Pt(12)
    mem3_title.paragraph_format.space_after = Pt(4)

    b7 = doc.add_paragraph(style=bullet_style)
    b7_run = b7.add_run("Xây dựng cấu trúc CSDL Quizzes: Thiết kế thành công 3 bảng dữ liệu quan hệ quizzes (đề thi), questions (câu hỏi liên kết qua quiz_id), và quiz_attempts (lượt làm bài thi của học viên). Thực thi kịch bản di chuyển migrate_quizzes.js dọn dẹp các bảng rác cũ và tạo dữ liệu trắc nghiệm mẫu hoàn chỉnh.")
    b7_run.font.name = 'Arial'
    b7_run.font.size = Pt(10)
    b7_run.font.color.rgb = CHARCOAL
    b7.paragraph_format.line_spacing = 1.15

    b8 = doc.add_paragraph(style=bullet_style)
    b8_run = b8.add_run("Xử lý lỗi lỗi Migration tự động (Vá lỗi Audit 3): Phối hợp bóc tách hoàn toàn các câu lệnh khởi tạo bảng DDL (CREATE TABLE) ra khỏi tệp database.js của backend. Đóng gói toàn bộ cấu trúc DB vào tệp schema.sql độc lập để Chương import thủ công vào SQL Editor của Supabase. Mã nguồn database.js hiện tại chỉ chạy health check kết nối, giải quyết triệt để lỗi deadlock khi khởi động máy chủ.")
    b8_run.font.name = 'Arial'
    b8_run.font.size = Pt(10)
    b8_run.font.color.rgb = CHARCOAL
    b8.paragraph_format.line_spacing = 1.15

    # ----------------------------------------------------
    # KẾT LUẬN
    # ----------------------------------------------------
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("III. KẾT LUẬN")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(13)
    h3_run.font.bold = True
    h3_run.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    p2_run = p2.add_run("Toàn bộ các thành viên trong nhóm phát triển dự án E-Learn Academy đã hoàn thành xuất sắc các mục tiêu đề ra cho tuần 5. Hệ thống hiện tại có cấu trúc cơ sở dữ liệu Supabase đồng bộ, backend xử lý API tối ưu hiệu năng và frontend kết nối dữ liệu động hoàn chỉnh.")
    p2_run.font.name = 'Arial'
    p2_run.font.size = Pt(10.5)
    p2_run.font.color.rgb = CHARCOAL
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(18)

    # Ký tên
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig_p.add_run("Hà Nội, ngày 02 tháng 07 năm 2026\nĐại diện nhóm thực hiện đồ án\n(Ký và ghi rõ họ tên)")
    sig_run.font.name = 'Arial'
    sig_run.font.size = Pt(10)
    sig_run.font.italic = True
    sig_run.font.color.rgb = CHARCOAL
    sig_p.paragraph_format.space_before = Pt(24)

    # Lưu file
    filename = "BAO CAO TIEN DO HOAN THANH TUAN 5.docx"
    doc.save(filename)
    print("Report generated successfully as: " + filename)

if __name__ == "__main__":
    create_report()
