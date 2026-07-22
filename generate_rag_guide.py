# -*- coding: utf-8 -*-
import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def make_callout_box(doc, text_list, title="LƯU Ý QUAN TRỌNG"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC") # slate-50
    set_cell_margins(cell, top=200, bottom=200, left=250, right=200)
    
    # Left border only (thick Smart Indigo blue)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders %s>'
        '<w:top w:val="none"/>'
        '<w:bottom w:val="none"/>'
        '<w:left w:val="single" w:sz="36" w:space="0" w:color="1D4ED8"/>'
        '<w:right w:val="none"/>'
        '</w:tcBorders>' % nsdecls('w')
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"★ {title}\n")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(29, 78, 216) # Smart Indigo
    
    for i, text in enumerate(text_list):
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate-600
        if i < len(text_list) - 1:
            p.add_run("\n")
            
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

def add_heading_1(doc, text, color):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(8)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = color
    return h

def add_heading_2(doc, text, color):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = color
    return h

def add_body_paragraph(doc, text, bold_words=[]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    
    current_idx = 0
    words_to_bold = sorted([(text.find(w), w) for w in bold_words if text.find(w) != -1], key=lambda x: x[0])
    
    for start_idx, word in words_to_bold:
        if start_idx < current_idx:
            continue
        # Add normal text before bold word
        if start_idx > current_idx:
            run_normal = p.add_run(text[current_idx:start_idx])
            run_normal.font.name = 'Times New Roman'
            run_normal.font.size = Pt(13)
        
        # Add bold word
        run_bold = p.add_run(word)
        run_bold.font.name = 'Times New Roman'
        run_bold.font.size = Pt(13)
        run_bold.font.bold = True
        
        current_idx = start_idx + len(word)
        
    if current_idx < len(text):
        run_end = p.add_run(text[current_idx:])
        run_end.font.name = 'Times New Roman'
        run_end.font.size = Pt(13)
        
    return p

def main():
    doc = Document()
    
    # Set page size to A4
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)     # 2.54 cm
        section.bottom_margin = Inches(1.0)  # 2.54 cm
        section.left_margin = Inches(1.18)   # 3.0 cm
        section.right_margin = Inches(0.78)  # 2.0 cm
        
    c_indigo = RGBColor(29, 78, 216) # #1d4ed8
    c_orange = RGBColor(255, 122, 48) # #ff7a30
    c_ink = RGBColor(15, 23, 42) # #0f172a
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.font.color.rgb = c_ink
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(6)
    
    # --- TITLE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(18)
    
    run_sub = p_title.add_run("BÁO CÁO PHÂN TÍCH KIẾN TRÚC\n")
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = c_orange
    
    run_main = p_title.add_run("CƠ CHẾ HOẠT ĐỘNG CỦA RAG CHATBOT AI\nVÀ TÍNH ỨNG DỤNG TRONG DỰ ÁN E-LEARNING E-LEARN ACADEMY\n")
    run_main.font.size = Pt(16)
    run_main.font.bold = True
    run_main.font.color.rgb = c_indigo
    
    run_time = p_title.add_run("(Tài liệu nghiên cứu kiến trúc hệ thống và hướng dẫn phát triển RAG)")
    run_time.font.size = Pt(11)
    run_time.font.italic = True
    run_time.font.color.rgb = RGBColor(100, 116, 139)
    
    # Metadata info box
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.left_indent = Inches(0.5)
    p_meta.paragraph_format.space_after = Pt(24)
    
    def add_meta_line(p, label, value):
        r_lbl = p.add_run(f"• {label}: ")
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(11)
        r_val = p.add_run(f"{value}\n")
        r_val.font.size = Pt(11)
        
    add_meta_line(p_meta, "Vai trò biên soạn", "Lập trình viên Phát triển Hệ thống (Developer)")
    add_meta_line(p_meta, "Đối tượng thụ hưởng", "Giảng viên (Lecturer) & Học viên (Student)")
    add_meta_line(p_meta, "Công nghệ RAG tích hợp", "NodeJS Express / Python Core Pipeline / Pinecone Vector Database / Gemini Embedding & LLM")
    
    # Horizontal Divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    r_div = p_div.add_run("_________________________________________________________________________________")
    r_div.font.color.rgb = RGBColor(226, 232, 240)
    r_div.font.bold = True
    
    # --- SECTION 1 ---
    add_heading_1(doc, "I. TỔNG QUAN VỀ CƠ CHẾ RAG (RETRIEVAL-AUGMENTED GENERATION)", c_indigo)
    
    add_body_paragraph(doc, 
        "RAG (Retrieval-Augmented Generation - Thế hệ tăng cường truy xuất) là một kỹ thuật tiên tiến trong kỹ nghệ AI, "
        "kết hợp sức mạnh của các Mô hình ngôn ngữ lớn (LLM - như Gemini, GPT) với một Hệ thống truy xuất thông tin độc lập. "
        "Mục tiêu cốt lõi của RAG là giải quyết hai nhược điểm lớn nhất của các mô hình LLM truyền thống: tính 'ảo tưởng' (hallucination) "
        "và việc thiếu khả năng truy cập thông tin nội bộ, riêng tư hoặc thông tin mới được cập nhật sau khi mô hình đã hoàn thành huấn luyện.",
        ["RAG", "ảo tưởng", "Gemini"]
    )
    
    add_body_paragraph(doc,
        "Thay vì hỏi LLM trực tiếp và để mô hình tự suy đoán câu trả lời dựa trên kho tri thức cũ có sẵn, cơ chế RAG hoạt động "
        "bằng cách chủ động đi tìm các đoạn tài liệu có liên quan nhất đến câu hỏi của người dùng từ một cơ sở dữ liệu tri thức đáng tin cậy. "
        "Sau đó, hệ thống ghép câu hỏi kèm các đoạn tài liệu tìm được thành một 'ngữ cảnh' (Context) gửi cho LLM. Nhờ đó, LLM chỉ đóng vai trò "
        "như một biên dịch viên trung gian, đọc hiểu dữ liệu được cung cấp và tổng hợp thành câu trả lời chính xác, trung thực 100% theo tài liệu.",
        ["ngữ cảnh", "cơ sở dữ liệu tri thức"]
    )
    
    make_callout_box(
        doc,
        [
            "Tránh ảo tưởng (Hallucination): Câu trả lời luôn dựa trên dữ liệu giảng dạy thực tế.",
            "Bảo mật và Cá nhân hóa: Không cần gửi dữ liệu nội bộ đi huấn luyện lại mô hình lớn.",
            "Tối ưu chi phí: Chi phí truy vấn Vector Database và Prompt nhỏ hơn hàng ngàn lần so với việc Fine-tuning (tinh chỉnh) lại LLM."
        ],
        "LỢI ÍCH CỐT LÕI CỦA RAG TRONG GIÁO DỤC TRỰC TUYẾN"
    )
    
    # --- SECTION 2 ---
    add_heading_1(doc, "II. CƠ CHẾ PHÂN LOẠI THEO BÀI HỌC (LESSON-SCOPED RAG)", c_indigo)
    
    add_body_paragraph(doc,
        "Trong một ứng dụng E-learning như E-Learn Academy, khóa học chứa rất nhiều bài giảng video và tài liệu PDF bổ trợ khác nhau. "
        "Nếu học viên đang xem bài giảng số 16 (Lesson 16) và hỏi về cấu trúc ngữ pháp có trong video, hệ thống RAG không thể tìm kiếm "
        "tất cả tài liệu của toàn bộ khóa học vì điều đó sẽ gây loãng thông tin và có nguy cơ lấy nhầm ngữ cảnh của bài học khác. "
        "Do đó, hệ thống triển khai cơ chế lọc theo bài học (Lesson-scoped RAG) dựa trên thuộc tính Metadata.",
        ["Lesson-scoped RAG", "Metadata"]
    )
    
    add_heading_2(doc, "1. Giai đoạn Phát triển (Môi trường Sandbox của Lập trình viên)", c_indigo)
    add_body_paragraph(doc,
        "Trong giai đoạn xây dựng hệ thống, Lập trình viên sử dụng một pipeline huấn luyện offline bằng ngôn ngữ Python để nạp dữ liệu thử nghiệm. "
        "Quá trình này được thực hiện thông qua script 'rag-training/main.py'. Tại đây, lập trình viên đặt các tài liệu text mẫu (ví dụ: 'lesson16-supplement.txt') "
        "vào thư mục 'rag-training/data'.",
        ["rag-training/main.py", "rag-training/data"]
    )
    add_body_paragraph(doc,
        "Khi khởi chạy, chương trình tự động trích xuất tên file thông qua biểu thức chính quy (Regular Expression) để tìm số bài học (ví dụ: lấy ra ID '16'). "
        "Sau đó, chương trình băm nhỏ văn bản (Chunking) thành nhiều khối, nhúng thành vector bằng mô hình nhúng của Gemini, và gắn thuộc tính metadata "
        "'lesson_id = 16' vào mỗi khối vector trước khi đẩy lên cơ sở dữ liệu Pinecone.",
        ["lesson_id = 16", "Pinecone"]
    )
    
    add_heading_2(doc, "2. Giai đoạn Vận hành thực tế (Production)", c_indigo)
    add_body_paragraph(doc,
        "Khi bàn giao sản phẩm thực tế, Giảng viên sẽ là người thao tác trực tiếp trên giao diện quản trị Web (Frontend) để tải lên tài liệu và video khóa học. "
        "Quy trình xử lý lúc này được tự động hóa hoàn toàn ở Backend NodeJS chứ không chạy thủ công qua script Python của lập trình viên:",
        ["Backend NodeJS"]
    )
    
    # Bullet points for production flow
    p_b1 = doc.add_paragraph(style='List Bullet')
    r = p_b1.add_run("Bước 1: ")
    r.bold = True
    p_b1.add_run("Giảng viên tải lên tài liệu học tập (PDF/TXT) hoặc video trực tiếp lên một Bài học cụ thể (ví dụ: Lesson ID = 25).")
    
    p_b2 = doc.add_paragraph(style='List Bullet')
    r = p_b2.add_run("Bước 2 (Xử lý Text & Audio): ")
    r.bold = True
    p_b2.add_run("Backend NodeJS nhận file. Nếu là PDF, server tự động trích xuất text. Nếu là video bài giảng, server gửi audio tới API chuyển đổi giọng nói thành văn bản (Speech-to-Text) để sinh ra transcript bài học.")
    
    p_b3 = doc.add_paragraph(style='List Bullet')
    r = p_b3.add_run("Bước 3 (Băm & Nhúng): ")
    r.bold = True
    p_b3.add_run("Đoạn văn bản trích xuất được cắt nhỏ và gửi đến Gemini Embedding API để chuyển thành vector đại diện ngữ nghĩa.")
    
    p_b4 = doc.add_paragraph(style='List Bullet')
    r = p_b4.add_run("Bước 4 (Đẩy lên VectorDB với Metadata): ")
    r.bold = True
    p_b4.add_run("Backend NodeJS lưu trữ trực tiếp các vector này vào Pinecone, tự động kèm nhãn metadata { lesson_id: 25 } tương ứng với khóa ngoại bài học trong database.")
    
    add_heading_2(doc, "3. Cơ chế Truy vấn Phản hồi Thông minh", c_indigo)
    add_body_paragraph(doc,
        "Khi học viên bấm vào nút 'Giải thích ngữ pháp trọng tâm' hoặc nhập câu hỏi ở thanh chat bài học: "
        "Frontend ReactJS sẽ gửi kèm nội dung câu hỏi cùng với tham số 'lessonId' hiện tại lên Backend. "
        "Tại đây, Backend NodeJS tạo vector truy vấn từ câu hỏi và thực hiện truy xuất dữ liệu từ Pinecone bằng cách thêm bộ lọc logic: "
        "queryOptions.filter = { lesson_id: { $eq: parsedLessonId } }.",
        ["queryOptions.filter", "lesson_id"]
    )
    add_body_paragraph(doc,
        "Nhờ bộ lọc này, Pinecone sẽ khoanh vùng tìm kiếm và loại bỏ 100% tài liệu của các bài học khác, chỉ trả về các đoạn ngữ cảnh "
        "thực sự thuộc về bài học mà học viên đang xem. Dữ liệu này sau đó được chuyển tiếp đến Gemini để tạo ra phản hồi chuyên sâu và chính xác nhất.",
        ["Pinecone", "Gemini"]
    )
    
    # --- SECTION 3 ---
    add_heading_1(doc, "III. TÍNH ỨNG DỤNG THỰC TIỄN CỦA RAG TRONG DỰ ÁN E-LEARN ACADEMY", c_indigo)
    
    add_body_paragraph(doc,
        "Việc tích hợp công nghệ RAG mang lại giá trị gia tăng cực kỳ lớn cho ứng dụng E-learning E-Learn Academy, "
        "nâng cấp trải nghiệm học tập truyền thống (chỉ xem và làm bài tập thụ động) thành trải nghiệm tương tác thông minh hai chiều chủ động.",
        ["E-Learn Academy"]
    )
    
    # Applications
    p_app1 = doc.add_paragraph(style='List Bullet')
    r = p_app1.add_run("1. Trợ lý ảo hỗ trợ học tập cá nhân hóa 24/7 (AI Learning Assistant):\n")
    r.bold = True
    p_app1.add_run("Học viên học trực tuyến thường gặp khó khăn do không thể tương tác trực tiếp với giảng viên khi có thắc mắc. RAG Chatbot đóng vai trò là một gia sư AI luôn sẵn sàng hỗ trợ, giải đáp ngay lập tức các thắc mắc về nội dung bài học, ngữ pháp hoặc từ vựng xuất hiện trong tài liệu phụ trợ.")
    
    p_app2 = doc.add_paragraph(style='List Bullet')
    r = p_app2.add_run("2. Tự động hóa giải đáp thắc mắc dựa trên video giảng dạy:\n")
    r.bold = True
    p_app2.add_run("Thay vì chỉ đọc tài liệu tĩnh, học viên có thể hỏi trực tiếp về các mốc thời gian hoặc nội dung giảng viên nói trong video (nhờ vào cơ chế nạp transcript video vào RAG). Học viên có thể hỏi các câu hỏi như 'Giảng viên đã nói gì về thì Hiện tại hoàn thành ở phút thứ 3?' và nhận câu trả lời chính xác.")
    
    p_app3 = doc.add_paragraph(style='List Bullet')
    r = p_app3.add_run("3. Giảm tải khối lượng công việc cho Giảng viên:\n")
    r.bold = True
    p_app3.add_run("Thông thường, giảng viên phải trả lời hàng trăm câu hỏi trùng lặp từ nhiều học viên khác nhau trên diễn đàn khóa học. Trợ lý ảo AI sẽ tự động xử lý và giải đáp đến 80-90% các câu hỏi thường gặp liên quan đến nội dung bài học, giúp giảng viên tập trung vào việc chuyên môn biên soạn bài giảng chất lượng cao hơn.")
    
    p_app4 = doc.add_paragraph(style='List Bullet')
    r = p_app4.add_run("4. Mô hình kinh doanh VIP Subscription hấp dẫn:\n")
    r.bold = True
    p_app4.add_run("Trong dự án E-Learn Academy, tính năng trò chuyện và hỏi đáp chuyên sâu với Trợ lý RAG Chatbot AI có thể được đóng gói làm một đặc quyền dành riêng cho học viên đăng ký tài khoản VIP. Đây là động lực thương mại hóa rất tốt, khuyến khích học viên nâng cấp gói trả phí để có trải nghiệm học tập tối ưu nhất.")
    
    # Final Callout
    make_callout_box(
        doc,
        [
            "Giúp lập trình viên nắm rõ luồng xử lý dữ liệu và cách tích hợp tự động hóa ở Backend.",
            "Làm tài liệu thuyết trình (Slides/Thesis Report) cực kỳ đắt giá cho Hội đồng bảo vệ đồ án tốt nghiệp.",
            "Khẳng định chiều sâu kỹ thuật của dự án E-Learn Academy so với các dự án E-learning thông thường khác."
        ],
        "Ý NGHĨA CỦA TÀI LIỆU PHÂN TÍCH RAG"
    )
    
    output_filename = "Bao_cao_co_che_hoat_dong_RAG.docx"
    doc.save(output_filename)
    print(f"✅ Báo cáo RAG đã được lưu thành công tại file: {output_filename}")

if __name__ == "__main__":
    main()
